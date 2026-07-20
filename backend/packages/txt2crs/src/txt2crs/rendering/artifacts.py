# SPDX-License-Identifier: MIT-0

"""Safe deterministic rendering for all learner and instructor artifacts."""

import re
import textwrap
import unicodedata
from dataclasses import dataclass
from html import escape
from io import BytesIO
from urllib.parse import urlsplit

import fitz  # type: ignore[import-untyped]
from docx import Document

from txt2crs.domain.models import (
    AnswerKey,
    Assessment,
    Course,
    ReviewPack,
)
from txt2crs.domain.validation import ArtifactBundle

_ACTIVE_CONTENT_PATTERNS = (
    re.compile(r"<\s*script\b", re.IGNORECASE),
    re.compile(r"<\s*iframe\b", re.IGNORECASE),
    re.compile(r"\bon[a-z]+\s*=", re.IGNORECASE),
    re.compile(
        r"<\s*(?:img|audio|video|source)\b[^>]*\bsrc\s*=\s*"
        r"['\"]?\s*https?://",
        re.IGNORECASE,
    ),
    re.compile(r"\bjavascript\s*:", re.IGNORECASE),
    re.compile(r"url\s*\(\s*['\"]?\s*https?://", re.IGNORECASE),
)
_PRIVATE_CONTENT_PATTERNS = (
    re.compile(r"(?i)\bBearer\s+[A-Za-z0-9._~+/-]+"),
    re.compile(r"(?:/home|/Users)/[^/\s]+/"),
    re.compile(r"(?i)\bsk-[A-Za-z0-9_-]{6,}"),
    re.compile(
        r"(?i)\b(?:api[_-]?key|access_token|refresh_token)\s*[:=]\s*"
        r"[A-Za-z0-9._~+/-]{6,}"
    ),
)
_INLINE_MARKDOWN_LINK_PATTERN = re.compile(r"\[([^\]\n]+)\]\((https?://[^)\s]+)\)")
_IDENTIFIER_CHARACTER_CLASS = r"A-Za-z0-9._:-"
_PDF_PUNCTUATION_TRANSLATION = str.maketrans(
    {
        "\u00a0": " ",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u2026": "...",
    }
)
_REVIEW_SCHEMA_FIELD_PATTERN = re.compile(
    r"`?\b(section|objective|exercise|flashcard|source)_id\b`?",
    re.IGNORECASE,
)
_BACKTICKED_REVIEW_IDENTIFIER_PATTERN = re.compile(
    r"`(?:practice|pe|worked|we|flashcard|fc|card|section|sec|objective|obj|lo)"
    r"[A-Za-z0-9._:-]+`",
    re.IGNORECASE,
)
_BARE_REVIEW_IDENTIFIER_PATTERN = re.compile(
    r"\b(?:section|objective|practice|exercise|worked|flashcard)"
    r"[-_:](?=[A-Za-z0-9._:-]*\d)[A-Za-z0-9._:-]+\b",
    re.IGNORECASE,
)
_BARE_SHORT_REVIEW_IDENTIFIER_PATTERN = re.compile(
    r"\b(?:lo|obj|sec|pe|we|fc)[-_:]?\d[A-Za-z0-9._:-]*\b",
    re.IGNORECASE,
)


class RenderedOutputQaError(ValueError):
    """Raised when deterministic post-render QA rejects an artifact."""


@dataclass(frozen=True, slots=True)
class RenderedArtifact:
    """One named byte artifact ready for private storage."""

    file_name: str
    media_type: str
    content: bytes


def derive_safe_filename(title: str) -> str:
    """Derive an ASCII path-safe slug without an AI naming call."""

    normalized_title = unicodedata.normalize("NFKD", title)
    ascii_title = normalized_title.encode("ascii", errors="ignore").decode("ascii")
    slug = re.sub(r"[^a-z0-9]+", "-", ascii_title.casefold()).strip("-")
    return slug[:100] or "course"


def validate_rendered_html(
    html_content: str,
    *,
    required_element_ids: set[str],
) -> None:
    """Reject missing structure, active content, remote media, and private data."""

    issues: list[str] = []
    for required_element_id in sorted(required_element_ids):
        identifier_pattern = re.compile(
            rf"\bid\s*=\s*['\"]{re.escape(required_element_id)}['\"]",
            re.IGNORECASE,
        )
        if identifier_pattern.search(html_content) is None:
            issues.append(f"missing required section: {required_element_id}")
    if any(pattern.search(html_content) for pattern in _ACTIVE_CONTENT_PATTERNS):
        issues.append("active or remote embedded content is forbidden")
    if any(pattern.search(html_content) for pattern in _PRIVATE_CONTENT_PATTERNS):
        issues.append("private diagnostic or credential-shaped content is forbidden")
    if issues:
        raise RenderedOutputQaError("; ".join(issues[:20]))


def _point_label(points: int) -> str:
    """Return grammatically correct learner-facing point text."""

    unit = "point" if points == 1 else "points"
    return f"{points} {unit}"


def _plain_text_from_inline_markdown(value: str) -> str:
    """Convert the renderer's small inline Markdown subset to portable text.

    HTML and Markdown retain their native link semantics. PDF and DOCX are
    intentionally dependency-light, so those formats spell out a link target
    and remove emphasis/code markers instead of exposing raw Markdown syntax.
    """

    text_with_readable_links = _INLINE_MARKDOWN_LINK_PATTERN.sub(
        r"\1 (\2)",
        value,
    )
    return text_with_readable_links.replace("**", "").replace("__", "").replace("`", "")


def _summary_text(value: str) -> str:
    """Remove one redundant leading colon after the fixed Summary label."""

    stripped_value = value.lstrip()
    if stripped_value.startswith(":"):
        return stripped_value.removeprefix(":").lstrip()
    return stripped_value


def _replace_identifier_references(
    value: str,
    *,
    display_label_by_identifier: dict[str, str],
) -> str:
    """Replace canonical validation IDs only when they occur as whole tokens."""

    reader_text = value
    # Longer identifiers run first so a short ID cannot consume part of a
    # longer, equally valid ID. Backtick-wrapped references are common in model
    # prose and are replaced as one unit so their Markdown markers disappear.
    for identifier in sorted(
        display_label_by_identifier,
        key=len,
        reverse=True,
    ):
        display_label = display_label_by_identifier[identifier]
        reader_text = reader_text.replace(f"`{identifier}`", display_label)
        reader_text = re.sub(
            rf"(?<![{_IDENTIFIER_CHARACTER_CLASS}])"
            rf"{re.escape(identifier)}"
            rf"(?![{_IDENTIFIER_CHARACTER_CLASS}])",
            display_label,
            reader_text,
        )
    return _humanize_unresolved_review_references(reader_text)


def _humanize_unresolved_review_references(value: str) -> str:
    """Remove stale identifier-shaped prose that has no canonical mapping."""

    reader_text = _REVIEW_SCHEMA_FIELD_PATTERN.sub(
        lambda match: match.group(1).casefold(),
        value,
    )

    def unresolved_identifier_label(match: re.Match[str]) -> str:
        """Return a neutral label without inventing a canonical target."""

        normalized_identifier = match.group(0).strip("`").casefold()
        if normalized_identifier.startswith(("practice", "pe-", "pe_")):
            return "Practice exercise"
        if normalized_identifier.startswith(("worked", "we-", "we_")):
            return "Worked example"
        if normalized_identifier.startswith(("flashcard", "card", "fc-", "fc_")):
            return "Flashcard"
        if normalized_identifier.startswith(("section", "sec-", "sec_")):
            return "section"
        return "objective"

    reader_text = _BACKTICKED_REVIEW_IDENTIFIER_PATTERN.sub(
        unresolved_identifier_label,
        reader_text,
    )

    def bare_identifier_label(match: re.Match[str]) -> str:
        """Return a contextual neutral label for an unquoted stale ID."""

        identifier_kind = re.split(r"[-_:]", match.group(0), maxsplit=1)[0].casefold()
        if identifier_kind == "section":
            return "the related section"
        if identifier_kind == "objective":
            return "the related objective"
        if identifier_kind in {"practice", "exercise"}:
            return "Practice exercise"
        if identifier_kind == "worked":
            return "Worked example"
        return "Flashcard"

    reader_text = _BARE_REVIEW_IDENTIFIER_PATTERN.sub(
        bare_identifier_label,
        reader_text,
    )

    def short_identifier_label(match: re.Match[str]) -> str:
        """Humanize compact IDs such as ``lo2`` or ``pe3``."""

        normalized_identifier = match.group(0).casefold()
        if normalized_identifier.startswith(("lo", "obj")):
            return "the related objective"
        if normalized_identifier.startswith("sec"):
            return "the related section"
        if normalized_identifier.startswith("pe"):
            return "Practice exercise"
        if normalized_identifier.startswith("we"):
            return "Worked example"
        return "Flashcard"

    reader_text = _BARE_SHORT_REVIEW_IDENTIFIER_PATTERN.sub(
        short_identifier_label,
        reader_text,
    )
    # A model may write "objective `lo1`". The known-ID replacement is
    # intentionally context-free, so collapse the resulting duplicated label
    # here after every canonical mapping is complete.
    return re.sub(
        r"\bobjective\s+Objective\s+(\d+)\b",
        r"Objective \1",
        reader_text,
        flags=re.IGNORECASE,
    )


class ArtifactRenderer:
    """Render one validated bundle to learner and instructor formats."""

    def render_bundle(
        self,
        bundle: ArtifactBundle,
    ) -> dict[str, RenderedArtifact]:
        """Return deterministic outputs and run QA before exposing any bytes."""

        base_name = derive_safe_filename(bundle.course.title)
        course_html = self._render_course_html(bundle.course)
        review_html = self._render_review_html(bundle.review_pack, bundle.course)
        assessment_html = self._render_assessment_html(
            bundle.assessment,
            bundle.course,
        )
        answer_key_html = self._render_answer_key_html(
            bundle.answer_key,
            bundle.assessment,
            bundle.course,
        )
        validate_rendered_html(
            course_html,
            required_element_ids={"course-content", "bibliography"},
        )
        validate_rendered_html(
            review_html,
            required_element_ids={"review-content"},
        )
        validate_rendered_html(
            assessment_html,
            required_element_ids={"assessment-content"},
        )
        validate_rendered_html(
            answer_key_html,
            required_element_ids={"answer-key-content"},
        )

        course_markdown = self._render_course_markdown(bundle.course)
        review_markdown = self._render_review_markdown(
            bundle.review_pack,
            bundle.course,
        )
        assessment_markdown = self._render_assessment_markdown(bundle.assessment)
        answer_key_markdown = self._render_answer_key_markdown(
            bundle.answer_key,
            bundle.assessment,
            bundle.course,
        )
        return {
            "course_html": RenderedArtifact(
                file_name=f"{base_name}-course.html",
                media_type="text/html; charset=utf-8",
                content=course_html.encode("utf-8"),
            ),
            "course_markdown": RenderedArtifact(
                file_name=f"{base_name}-course.md",
                media_type="text/markdown; charset=utf-8",
                content=course_markdown.encode("utf-8"),
            ),
            "course_pdf": RenderedArtifact(
                file_name=f"{base_name}-course.pdf",
                media_type="application/pdf",
                content=self._render_course_pdf(bundle.course),
            ),
            "course_docx": RenderedArtifact(
                file_name=f"{base_name}-course.docx",
                media_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                content=self._render_docx(
                    title=bundle.course.title,
                    subject="txt2crs generated course",
                    markdown_content=course_markdown,
                ),
            ),
            "review_pack_html": RenderedArtifact(
                file_name=f"{base_name}-review-pack.html",
                media_type="text/html; charset=utf-8",
                content=review_html.encode("utf-8"),
            ),
            "review_pack_markdown": RenderedArtifact(
                file_name=f"{base_name}-review-pack.md",
                media_type="text/markdown; charset=utf-8",
                content=review_markdown.encode("utf-8"),
            ),
            "review_pack_pdf": RenderedArtifact(
                file_name=f"{base_name}-review-pack.pdf",
                media_type="application/pdf",
                content=self._render_text_pdf(
                    title=f"{bundle.course.title} Review Pack",
                    subject="txt2crs generated review pack",
                    lines=review_markdown.splitlines(),
                ),
            ),
            "review_pack_docx": RenderedArtifact(
                file_name=f"{base_name}-review-pack.docx",
                media_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                content=self._render_docx(
                    title=f"{bundle.course.title} Review Pack",
                    subject="txt2crs generated review pack",
                    markdown_content=review_markdown,
                ),
            ),
            "assessment_html": RenderedArtifact(
                file_name=f"{base_name}-assessment.html",
                media_type="text/html; charset=utf-8",
                content=assessment_html.encode("utf-8"),
            ),
            "assessment_markdown": RenderedArtifact(
                file_name=f"{base_name}-assessment.md",
                media_type="text/markdown; charset=utf-8",
                content=assessment_markdown.encode("utf-8"),
            ),
            "assessment_pdf": RenderedArtifact(
                file_name=f"{base_name}-assessment.pdf",
                media_type="application/pdf",
                content=self._render_text_pdf(
                    title=bundle.assessment.title,
                    subject="txt2crs generated student assessment",
                    lines=assessment_markdown.splitlines(),
                ),
            ),
            "assessment_docx": RenderedArtifact(
                file_name=f"{base_name}-assessment.docx",
                media_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                content=self._render_docx(
                    title=bundle.assessment.title,
                    subject="txt2crs generated student assessment",
                    markdown_content=assessment_markdown,
                ),
            ),
            "answer_key_html": RenderedArtifact(
                file_name=f"{base_name}-answer-key.html",
                media_type="text/html; charset=utf-8",
                content=answer_key_html.encode("utf-8"),
            ),
            "answer_key_markdown": RenderedArtifact(
                file_name=f"{base_name}-answer-key.md",
                media_type="text/markdown; charset=utf-8",
                content=answer_key_markdown.encode("utf-8"),
            ),
            "answer_key_pdf": RenderedArtifact(
                file_name=f"{base_name}-answer-key.pdf",
                media_type="application/pdf",
                content=self._render_text_pdf(
                    title=f"{bundle.assessment.title} Answer Key",
                    subject="txt2crs generated instructor answer key",
                    lines=answer_key_markdown.splitlines(),
                ),
            ),
            "answer_key_docx": RenderedArtifact(
                file_name=f"{base_name}-answer-key.docx",
                media_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                content=self._render_docx(
                    title=f"{bundle.assessment.title} Answer Key",
                    subject="txt2crs generated instructor answer key",
                    markdown_content=answer_key_markdown,
                ),
            ),
        }

    def _render_docx(
        self,
        *,
        title: str,
        subject: str,
        markdown_content: str,
    ) -> bytes:
        """Create a real Word document from trusted deterministic Markdown.

        The DOCX renderer intentionally supports the small Markdown vocabulary
        emitted by this module. Keeping the conversion local means no model or
        remote document service can silently rewrite the accepted artifact.
        """

        document = Document()
        document.core_properties.title = title
        document.core_properties.subject = subject
        document.core_properties.author = "txt2crs"

        inside_code_block = False
        for raw_line in markdown_content.splitlines():
            stripped_line = raw_line.strip()
            if stripped_line == "```":
                inside_code_block = not inside_code_block
                continue
            if not stripped_line:
                document.add_paragraph()
                continue
            if inside_code_block:
                document.add_paragraph(raw_line, style="No Spacing")
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped_line)
            if heading_match is not None:
                heading_level = min(9, len(heading_match.group(1)))
                document.add_heading(
                    _plain_text_from_inline_markdown(heading_match.group(2)),
                    level=heading_level,
                )
                continue
            if stripped_line.startswith("- "):
                document.add_paragraph(
                    _plain_text_from_inline_markdown(stripped_line.removeprefix("- ")),
                    style="List Bullet",
                )
                continue
            numbered_match = re.match(r"^\d+\.\s+(.+)$", stripped_line)
            if numbered_match is not None:
                document.add_paragraph(
                    _plain_text_from_inline_markdown(numbered_match.group(1)),
                    style="List Number",
                )
                continue
            document.add_paragraph(_plain_text_from_inline_markdown(stripped_line))

        output_buffer = BytesIO()
        document.save(output_buffer)
        return output_buffer.getvalue()

    def _document_start(self, *, title: str, language: str) -> str:
        """Return shared accessible HTML document metadata."""

        direction = "rtl" if language in {"ar", "fa", "he", "ur"} else "ltr"
        return (
            "<!doctype html>"
            f'<html lang="{escape(language)}" dir="{direction}">'
            '<head><meta charset="utf-8">'
            '<meta name="viewport" content="width=device-width,initial-scale=1">'
            f"<title>{escape(title)}</title></head><body>"
        )

    def _render_course_html(self, course: Course) -> str:
        """Render semantic headings, course claims, and a bibliography."""

        evidence_by_id = {
            evidence.evidence_id: evidence for evidence in course.evidence
        }
        citation_by_location = {
            citation.artifact_location: citation for citation in course.citations
        }
        source_number = {
            source.source_id: index
            for index, source in enumerate(course.sources, start=1)
        }
        html_parts = [
            self._document_start(title=course.title, language=course.language),
            '<main id="course-content">',
            f"<h1>{escape(course.title)}</h1>",
            f"<p><strong>Audience:</strong> {escape(course.audience)}</p>",
            f"<p><strong>Level:</strong> {escape(course.level)}</p>",
            "<section><h2>Prerequisites</h2><ul>",
        ]
        html_parts.extend(
            f"<li>{escape(prerequisite)}</li>" for prerequisite in course.prerequisites
        )
        html_parts.extend(
            [
                "</ul></section>",
                '<section aria-labelledby="objectives-heading">',
                '<h2 id="objectives-heading">Learning objectives</h2><ul>',
            ]
        )
        html_parts.extend(
            f"<li>{escape(objective.description)}</li>"
            for objective in course.learning_objectives
        )
        html_parts.append("</ul></section>")
        for course_module in course.modules:
            html_parts.append(
                f'<section id="{escape(course_module.module_id)}">'
                f"<h2>{escape(course_module.title)}</h2>"
                f"<p>{escape(course_module.summary)}</p>"
            )
            for section in course_module.sections:
                html_parts.append(
                    f'<section id="{escape(section.section_id)}">'
                    f"<h3>{escape(section.title)}</h3>"
                )
                for block in section.content_blocks:
                    rendered_text = escape(block.text)
                    if block.kind == "code":
                        rendered_block = f"<pre><code>{rendered_text}</code></pre>"
                    else:
                        rendered_block = f"<p>{rendered_text}</p>"
                    citation = citation_by_location.get(block.block_id)
                    if citation is not None:
                        citation_numbers = sorted(
                            {
                                source_number[evidence_by_id[evidence_id].source_id]
                                for evidence_id in citation.evidence_ids
                            }
                        )
                        rendered_block += "".join(
                            f'<sup><a href="#source-{number}" '
                            f'aria-label="Source {number}">[{number}]</a></sup>'
                            for number in citation_numbers
                        )
                    html_parts.append(rendered_block)
                html_parts.append(
                    f"<p><strong>Summary:</strong> "
                    f"{escape(_summary_text(section.summary))}</p>"
                    "</section>"
                )
            if course_module.misconceptions:
                html_parts.append("<h3>Common misconceptions</h3><ul>")
                html_parts.extend(
                    f"<li>{escape(misconception)}</li>"
                    for misconception in course_module.misconceptions
                )
                html_parts.append("</ul>")
            if course_module.examples:
                html_parts.append("<h3>Examples</h3><ul>")
                html_parts.extend(
                    f"<li>{escape(example)}</li>" for example in course_module.examples
                )
                html_parts.append("</ul>")
            html_parts.append("</section>")
        html_parts.append("<section><h2>Glossary</h2><dl>")
        for glossary_term in course.glossary:
            html_parts.append(
                f"<dt>{escape(glossary_term.term)}</dt>"
                f"<dd>{escape(glossary_term.definition)}</dd>"
            )
        html_parts.append("</dl></section>")
        if course.unresolved_or_conflicting_claims:
            html_parts.append("<section><h2>Unresolved or conflicting claims</h2><ul>")
            html_parts.extend(
                f"<li>{escape(claim)}</li>"
                for claim in course.unresolved_or_conflicting_claims
            )
            html_parts.append("</ul></section>")
        html_parts.append('</main><section id="bibliography"><h2>Sources</h2><ol>')
        for index, source in enumerate(course.sources, start=1):
            title = escape(source.title)
            href = _safe_http_href(source.canonical_url)
            source_text = (
                f'<a href="{escape(href, quote=True)}" '
                'rel="noopener noreferrer">'
                f"{title}</a>"
                if href is not None
                else title
            )
            html_parts.append(
                f'<li id="source-{index}">{source_text} — '
                f"{escape(source.publisher_or_author)}</li>"
            )
        html_parts.append("</ol></section></body></html>")
        return "".join(html_parts)

    def _render_review_html(self, review: ReviewPack, course: Course) -> str:
        """Render the comprehensive review pack."""

        display_label_by_identifier = self._review_display_labels(
            review=review,
            course=course,
        )
        parts = [
            self._document_start(
                title=f"{course.title} Review Pack",
                language=course.language,
            ),
            '<main id="review-content">',
            f"<h1>{escape(course.title)} Review Pack</h1>",
            "<section><h2>Suggested review sequence</h2><ol>",
        ]
        parts.extend(
            "<li>"
            + escape(
                _replace_identifier_references(
                    review_step,
                    display_label_by_identifier=display_label_by_identifier,
                )
            )
            + "</li>"
            for review_step in review.review_sequence
        )
        parts.extend(
            [
                "</ol></section>",
                "<section><h2>Study guide</h2>",
            ]
        )
        source_by_id = {source.source_id: source for source in course.sources}
        for guide_item in review.study_guide:
            objective = next(
                objective
                for objective in course.learning_objectives
                if objective.objective_id == guide_item.objective_id
            )
            parts.append(
                "<article><h3>"
                f"{escape(display_label_by_identifier[guide_item.objective_id])}: "
                f"{escape(objective.description)}</h3>"
                f"<p>{escape(guide_item.summary)}</p><ul>"
            )
            parts.extend(
                f"<li>{escape(takeaway)}</li>" for takeaway in guide_item.key_takeaways
            )
            parts.append("</ul><h4>Common misconceptions</h4><ul>")
            parts.extend(
                f"<li>{escape(misconception)}</li>"
                for misconception in guide_item.misconceptions
            )
            parts.append("</ul><p><strong>Sources:</strong> ")
            source_links: list[str] = []
            for source_id in guide_item.source_ids:
                source = source_by_id[source_id]
                safe_href = _safe_http_href(source.canonical_url)
                if safe_href is None:
                    source_links.append(escape(source.title))
                else:
                    source_links.append(
                        f'<a href="{escape(safe_href, quote=True)}" '
                        'rel="noopener noreferrer">'
                        f"{escape(source.title)}</a>"
                    )
            parts.append(", ".join(source_links) or "Course content")
            parts.append("</p></article>")
        parts.append("</section><section><h2>Glossary</h2><dl>")
        for glossary_term in review.glossary:
            parts.append(
                f"<dt>{escape(glossary_term.term)}</dt>"
                f"<dd>{escape(glossary_term.definition)}</dd>"
            )
        parts.append("</section><section><h2>Flashcards</h2><dl>")
        for flashcard in review.flashcards:
            parts.append(
                f"<dt>{escape(flashcard.prompt)}</dt>"
                f"<dd>{escape(flashcard.answer)}</dd>"
            )
        parts.append("</dl></section><section><h2>Worked examples</h2>")
        for example_index, example in enumerate(review.worked_examples, start=1):
            parts.append(
                f"<article><h3>Worked example {example_index}</h3>"
                f"<p>{escape(example.prompt)}</p>"
                f"<p><strong>Worked solution:</strong> "
                f"{escape(example.solution)}</p></article>"
            )
        parts.append("</section><section><h2>Practice</h2>")
        for exercise_index, exercise in enumerate(
            review.practice_exercises,
            start=1,
        ):
            parts.append(
                f"<article><h3>Practice exercise {exercise_index}</h3>"
                f"<p>{escape(exercise.prompt)}</p>"
                f"<details><summary>Solution</summary>"
                f"<p>{escape(exercise.solution)}</p></details></article>"
            )
        parts.append("</section><section><h2>Section summaries</h2><dl>")
        for section_id, summary in sorted(review.section_summaries.items()):
            parts.append(
                f"<dt>{escape(display_label_by_identifier[section_id])}</dt>"
                f"<dd>{escape(summary)}</dd>"
            )
        parts.append(
            f"</dl></section><section><h2>Cumulative summary</h2>"
            f"<p>{escape(review.cumulative_summary)}</p></section>"
            "</main></body></html>"
        )
        return "".join(parts)

    def _review_display_labels(
        self,
        *,
        review: ReviewPack,
        course: Course,
    ) -> dict[str, str]:
        """Map internal cross-artifact IDs to stable reader-facing labels."""

        display_labels = {
            objective.objective_id: f"Objective {objective_index}"
            for objective_index, objective in enumerate(
                course.learning_objectives,
                start=1,
            )
        }
        display_labels.update(
            {
                course_module.module_id: course_module.title
                for course_module in course.modules
            }
        )
        display_labels.update(
            {
                section.section_id: section.title
                for course_module in course.modules
                for section in course_module.sections
            }
        )
        display_labels.update(
            {
                flashcard.flashcard_id: f"Flashcard {flashcard_index}"
                for flashcard_index, flashcard in enumerate(
                    review.flashcards,
                    start=1,
                )
            }
        )
        display_labels.update(
            {
                example.exercise_id: f"Worked example {example_index}"
                for example_index, example in enumerate(
                    review.worked_examples,
                    start=1,
                )
            }
        )
        display_labels.update(
            {
                exercise.exercise_id: f"Practice exercise {exercise_index}"
                for exercise_index, exercise in enumerate(
                    review.practice_exercises,
                    start=1,
                )
            }
        )
        return display_labels

    def _render_assessment_html(
        self,
        assessment: Assessment,
        course: Course,
    ) -> str:
        """Render only student-visible questions and point values."""

        parts = [
            self._document_start(
                title=assessment.title,
                language=course.language,
            ),
            '<main id="assessment-content">',
            f"<h1>{escape(assessment.title)}</h1>",
            f"<p>{escape(assessment.instructions)}</p>",
            f"<p><strong>Passing score:</strong> {assessment.passing_percentage}%</p>",
            "<ol>",
        ]
        for item in assessment.items:
            parts.append(
                f'<li id="{escape(item.item_id)}">'
                f"<p>{escape(item.prompt)} ({_point_label(item.points)})</p>"
            )
            if item.options:
                parts.append('<ol type="A">')
                parts.extend(f"<li>{escape(option)}</li>" for option in item.options)
                parts.append("</ol>")
            parts.append("</li>")
        parts.append("</ol></main></body></html>")
        return "".join(parts)

    def _render_answer_key_html(
        self,
        answer_key: AnswerKey,
        assessment: Assessment,
        course: Course,
    ) -> str:
        """Render instructor-only answers, explanations, and rubrics."""

        item_by_id = {item.item_id: item for item in assessment.items}
        evidence_by_id = {
            evidence.evidence_id: evidence for evidence in course.evidence
        }
        source_by_id = {source.source_id: source for source in course.sources}
        parts = [
            self._document_start(
                title=f"{assessment.title} Answer Key",
                language=course.language,
            ),
            '<main id="answer-key-content">',
            f"<h1>{escape(assessment.title)} Answer Key</h1>",
        ]
        for answer in answer_key.answers:
            item = item_by_id[answer.item_id]
            answer_source_ids = list(
                dict.fromkeys(
                    evidence_by_id[evidence_id].source_id
                    for evidence_id in answer.evidence_ids
                )
            )
            parts.append(
                f'<article id="answer-{escape(answer.item_id)}">'
                f"<h2>{escape(item.prompt)}</h2>"
                f"<p><strong>Answer:</strong> "
                f"{escape('; '.join(answer.correct_answers))}</p>"
                f"<p>{escape(answer.explanation)}</p>"
                "<h3>Evidence sources</h3><ul>"
            )
            for source_id in answer_source_ids:
                source = source_by_id[source_id]
                safe_href = _safe_http_href(source.canonical_url)
                source_text = escape(source.title)
                if safe_href is not None:
                    source_text = (
                        f'<a href="{escape(safe_href, quote=True)}" '
                        'rel="noopener noreferrer">'
                        f"{source_text}</a>"
                    )
                parts.append(f"<li>{source_text}</li>")
            parts.append("</ul><h3>Grading criteria</h3><ul>")
            parts.extend(
                f"<li>{escape(grading_criterion)}</li>"
                for grading_criterion in answer.grading_criteria
            )
            parts.append("</ul><h3>Rubric</h3><ul>")
            parts.extend(
                f"<li>{escape(criterion.criterion)}: "
                f"{criterion.points} — {escape(criterion.description)}</li>"
                for criterion in answer.rubric
            )
            parts.append("</ul></article>")
        parts.append("</main></body></html>")
        return "".join(parts)

    def _render_course_markdown(self, course: Course) -> str:
        """Render readable Markdown from the canonical course."""

        lines = [
            f"# {course.title}",
            "",
            f"**Audience:** {course.audience}",
            f"**Level:** {course.level}",
            "",
            "## Prerequisites",
            "",
        ]
        lines.extend(f"- {prerequisite}" for prerequisite in course.prerequisites)
        lines.extend(["", "## Learning objectives", ""])
        lines.extend(
            f"- {objective.description}" for objective in course.learning_objectives
        )
        lines.append("")
        for course_module in course.modules:
            lines.extend(
                [
                    f"## {course_module.title}",
                    "",
                    course_module.summary,
                    "",
                ]
            )
            for section in course_module.sections:
                lines.extend([f"### {section.title}", ""])
                lines.extend(block.text for block in section.content_blocks)
                lines.extend(["", f"**Summary:** {_summary_text(section.summary)}", ""])
            if course_module.misconceptions:
                lines.extend(["### Common misconceptions", ""])
                lines.extend(
                    f"- {misconception}"
                    for misconception in course_module.misconceptions
                )
                lines.append("")
            if course_module.examples:
                lines.extend(["### Examples", ""])
                lines.extend(f"- {example}" for example in course_module.examples)
            lines.append("")
        lines.extend(["## Glossary", ""])
        lines.extend(
            f"- **{glossary_term.term}:** {glossary_term.definition}"
            for glossary_term in course.glossary
        )
        if course.unresolved_or_conflicting_claims:
            lines.extend(["", "## Unresolved or conflicting claims", ""])
            lines.extend(
                f"- {claim}" for claim in course.unresolved_or_conflicting_claims
            )
        lines.extend(["## Sources", ""])
        lines.extend(
            f"- [{source.title}]({source.canonical_url})" for source in course.sources
        )
        return "\n".join(lines).strip() + "\n"

    def _render_review_markdown(
        self,
        review: ReviewPack,
        course: Course,
    ) -> str:
        """Render every canonical review component in a portable text format."""

        source_by_id = {source.source_id: source for source in course.sources}
        display_label_by_identifier = self._review_display_labels(
            review=review,
            course=course,
        )
        lines = [
            f"# {course.title} Review Pack",
            "",
            "## Suggested review sequence",
            "",
        ]
        lines.extend(
            f"{step_number}. "
            + _replace_identifier_references(
                step,
                display_label_by_identifier=display_label_by_identifier,
            )
            for step_number, step in enumerate(review.review_sequence, start=1)
        )
        lines.extend(["", "## Study guide", ""])
        objective_by_id = {
            objective.objective_id: objective
            for objective in course.learning_objectives
        }
        for guide_item in review.study_guide:
            lines.extend(
                [
                    "### "
                    f"{display_label_by_identifier[guide_item.objective_id]}: "
                    f"{objective_by_id[guide_item.objective_id].description}",
                    "",
                    guide_item.summary,
                    "",
                    "**Key takeaways**",
                    "",
                ]
            )
            lines.extend(f"- {takeaway}" for takeaway in guide_item.key_takeaways)
            lines.extend(["", "**Common misconceptions**", ""])
            lines.extend(
                f"- {misconception}" for misconception in guide_item.misconceptions
            )
            lines.extend(["", "**Sources**", ""])
            lines.extend(
                f"- [{source_by_id[source_id].title}]"
                f"({source_by_id[source_id].canonical_url})"
                for source_id in guide_item.source_ids
            )
            lines.append("")
        lines.extend(["## Glossary", ""])
        lines.extend(
            f"- **{term.term}:** {term.definition}" for term in review.glossary
        )
        lines.extend(["", "## Flashcards", ""])
        for flashcard in review.flashcards:
            lines.extend(
                [
                    f"### {flashcard.prompt}",
                    "",
                    flashcard.answer,
                    "",
                ]
            )
        lines.extend(["## Worked examples", ""])
        for example_index, example in enumerate(review.worked_examples, start=1):
            lines.extend(
                [
                    f"### Worked example {example_index}",
                    "",
                    example.prompt,
                    "",
                    f"**Worked solution:** {example.solution}",
                    "",
                ]
            )
        lines.extend(["## Practice exercises", ""])
        for exercise_index, exercise in enumerate(
            review.practice_exercises,
            start=1,
        ):
            lines.extend(
                [
                    f"### Practice exercise {exercise_index}",
                    "",
                    exercise.prompt,
                    "",
                    f"**Solution:** {exercise.solution}",
                    "",
                ]
            )
        lines.extend(["## Section summaries", ""])
        lines.extend(
            f"- **{display_label_by_identifier[section_id]}:** {summary}"
            for section_id, summary in sorted(review.section_summaries.items())
        )
        lines.extend(
            [
                "",
                "## Cumulative summary",
                "",
                review.cumulative_summary,
                "",
            ]
        )
        return "\n".join(lines).strip() + "\n"

    def _render_assessment_markdown(self, assessment: Assessment) -> str:
        """Render the student form without instructor-only answers."""

        lines = [
            f"# {assessment.title}",
            "",
            assessment.instructions,
            "",
            f"**Passing score:** {assessment.passing_percentage}%",
            "",
        ]
        for item_number, item in enumerate(assessment.items, start=1):
            lines.extend(
                [
                    f"## {item_number}. {item.prompt} ({_point_label(item.points)})",
                    "",
                ]
            )
            lines.extend(
                f"{chr(65 + option_index)}. {option}"
                for option_index, option in enumerate(item.options)
            )
            lines.append("")
        return "\n".join(lines).strip() + "\n"

    def _render_answer_key_markdown(
        self,
        answer_key: AnswerKey,
        assessment: Assessment,
        course: Course,
    ) -> str:
        """Render answers, reasoning, criteria, and point-based rubrics."""

        item_by_id = {item.item_id: item for item in assessment.items}
        evidence_by_id = {
            evidence.evidence_id: evidence for evidence in course.evidence
        }
        source_by_id = {source.source_id: source for source in course.sources}
        lines = [f"# {assessment.title} Answer Key", ""]
        for answer in answer_key.answers:
            item = item_by_id[answer.item_id]
            answer_source_ids = list(
                dict.fromkeys(
                    evidence_by_id[evidence_id].source_id
                    for evidence_id in answer.evidence_ids
                )
            )
            lines.extend(
                [
                    f"## {item.prompt}",
                    "",
                    f"**Answer:** {'; '.join(answer.correct_answers)}",
                    "",
                    answer.explanation,
                    "",
                    "**Evidence sources**",
                    "",
                ]
            )
            lines.extend(
                f"- [{source_by_id[source_id].title}]"
                f"({source_by_id[source_id].canonical_url})"
                for source_id in answer_source_ids
            )
            lines.extend(
                [
                    "",
                    "**Grading criteria**",
                    "",
                ]
            )
            lines.extend(
                f"- {grading_criterion}"
                for grading_criterion in answer.grading_criteria
            )
            lines.extend(["", "**Rubric**", ""])
            lines.extend(
                f"- {criterion.criterion}: {criterion.points} — {criterion.description}"
                for criterion in answer.rubric
            )
            lines.append("")
        return "\n".join(lines).strip() + "\n"

    def _render_course_pdf(self, course: Course) -> bytes:
        """Create a searchable text PDF without remote assets or AI formatting."""

        return self._render_text_pdf(
            title=course.title,
            subject="txt2crs generated course",
            lines=self._render_course_markdown(course).splitlines(),
        )

    def _render_text_pdf(
        self,
        *,
        title: str,
        subject: str,
        lines: list[str],
    ) -> bytes:
        """Create a searchable PDF from already validated local text lines."""

        document = fitz.open()
        try:
            page = document.new_page()
            vertical_position = 72.0
            for original_line in lines:
                # Markdown markers add no value in PDF output. Removing the
                # small safe subset keeps the text readable without parsing or
                # rendering arbitrary HTML.
                plain_line = re.sub(
                    r"^(?:#{1,6}\s+|[-*]\s+|\d+\.\s+)", "", original_line
                )
                plain_line = _plain_text_from_inline_markdown(plain_line)
                plain_line = plain_line.translate(_PDF_PUNCTUATION_TRANSLATION)
                wrapped_lines = textwrap.wrap(plain_line, width=85) or [""]
                for wrapped_line in wrapped_lines:
                    if vertical_position > 770:
                        page = document.new_page()
                        vertical_position = 72.0
                    # Helvetica creates searchable text. Unsupported glyphs are
                    # replaced by the PDF engine, while the HTML remains the
                    # fully accessible multilingual source of truth.
                    page.insert_text(
                        (72, vertical_position),
                        wrapped_line,
                        fontname="helv",
                        fontsize=11,
                    )
                    vertical_position += 16
            document.set_metadata(
                {
                    "title": title,
                    "subject": subject,
                    "author": "txt2crs",
                }
            )
            return bytes(document.tobytes(garbage=4, deflate=True))
        finally:
            document.close()


def _safe_http_href(url: str) -> str | None:
    """Allow bibliography links only for plain public HTTP(S) URL shapes."""

    parsed_url = urlsplit(url)
    if (
        parsed_url.scheme not in {"http", "https"}
        or parsed_url.hostname is None
        or parsed_url.username is not None
        or parsed_url.password is not None
    ):
        return None
    return url
