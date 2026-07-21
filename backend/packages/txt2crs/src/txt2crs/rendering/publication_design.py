# SPDX-License-Identifier: MIT-0

"""Shared editorial design primitives for deterministic publications.

The engine emits four publications in several file formats. This module keeps
their visual language consistent without making one format depend on another.
No remote font, image, stylesheet, conversion service, or model call
participates in rendering.
"""

import re
import textwrap
from dataclasses import dataclass
from io import BytesIO
from typing import Final, Literal

import fitz  # type: ignore[import-untyped]
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

PublicationKind = Literal["course", "review", "assessment", "answer-key"]

_INLINE_MARKDOWN_LINK_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"\[([^\]\n]+)\]\((https?://[^)\s]+)\)"
)
_INLINE_DOCX_TOKEN_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"(\*\*[^*\n]+\*\*|`[^`\n]+`|\[[^\]\n]+\]\(https?://[^)\s]+\))"
)
_PDF_PUNCTUATION_TRANSLATION: Final[dict[int, str]] = str.maketrans(
    {
        "\u00a0": " ",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u2026": "...",
    }
)

# These colors mirror the warm editorial palette in frontend/src/index.css.
_FOREST_HEX: Final[str] = "1A5038"
_FOREST_RGB: Final[tuple[float, float, float]] = (0.102, 0.314, 0.220)
_GOLD_RGB: Final[tuple[float, float, float]] = (0.722, 0.514, 0.165)
_INK_RGB: Final[tuple[float, float, float]] = (0.125, 0.118, 0.102)
_MUTED_RGB: Final[tuple[float, float, float]] = (0.410, 0.392, 0.350)
_RULE_RGB: Final[tuple[float, float, float]] = (0.835, 0.808, 0.745)


@dataclass(frozen=True, slots=True)
class PublicationTheme:
    """One named publication treatment inside the shared visual system."""

    label: str
    deck: str
    accent_hex: str
    accent_rgb: tuple[float, float, float]


_PUBLICATION_THEMES: Final[dict[PublicationKind, PublicationTheme]] = {
    "course": PublicationTheme(
        label="Course",
        deck="A research-grounded learning journey, designed for focused study.",
        accent_hex="2D6B4A",
        accent_rgb=(0.176, 0.420, 0.290),
    ),
    "review": PublicationTheme(
        label="Review pack",
        deck="Recall, connect, and practice the ideas that matter most.",
        accent_hex="3E6485",
        accent_rgb=(0.243, 0.392, 0.522),
    ),
    "assessment": PublicationTheme(
        label="Student assessment",
        deck="A clear, printable checkpoint for demonstrating understanding.",
        accent_hex="A06D18",
        accent_rgb=(0.627, 0.427, 0.094),
    ),
    "answer-key": PublicationTheme(
        label="Instructor answer key",
        deck="Instructor edition with evidence, grading guidance, and rubrics.",
        accent_hex="8A4738",
        accent_rgb=(0.541, 0.278, 0.220),
    ),
}


# The stylesheet is intentionally self-contained so downloaded HTML remains
# beautiful offline and preview sanitization does not need a network origin.
PUBLICATION_CSS: Final[str] = """
:root {
  color-scheme: light;
  --paper: #fbf8f0;
  --paper-raised: #fffdf8;
  --ink: #211f1b;
  --muted: #69645a;
  --forest: #1a5038;
  --gold: #b8832a;
  --rule: #ded7c8;
  --publication-accent: #2d6b4a;
  --display: Georgia, "Times New Roman", serif;
  --body: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
  --mono: ui-monospace, "SFMono-Regular", Consolas, "Liberation Mono", monospace;
}
* { box-sizing: border-box; }
html { background: #e9e4d8; scroll-behavior: smooth; }
body.publication-shell {
  margin: 0;
  color: var(--ink);
  background: var(--paper);
  font-family: var(--body);
  font-size: 1rem;
  line-height: 1.68;
  overflow-wrap: anywhere;
  text-rendering: optimizeLegibility;
}
body.publication--review { --publication-accent: #3e6485; }
body.publication--assessment { --publication-accent: #a06d18; }
body.publication--answer-key { --publication-accent: #8a4738; }
.publication-hero {
  position: relative;
  overflow: hidden;
  min-height: 25rem;
  padding: clamp(2rem, 7vw, 5rem) max(1.5rem, calc((100vw - 58rem) / 2));
  color: #fffdf8;
  background: var(--forest);
  border-top: 0.45rem solid var(--publication-accent);
  display: flex;
  flex-direction: column;
  justify-content: center;
}
.publication-hero::after {
  content: "";
  position: absolute;
  inset: auto -5rem -9rem auto;
  width: 22rem;
  height: 22rem;
  border: 1px solid rgb(255 255 255 / 0.18);
  border-radius: 50%;
  box-shadow: 0 0 0 3.25rem rgb(255 255 255 / 0.035),
    0 0 0 6.5rem rgb(255 255 255 / 0.025);
}
.publication-brand {
  position: relative;
  z-index: 1;
  display: flex;
  align-items: center;
  gap: 0.7rem;
  margin-bottom: clamp(3.5rem, 8vw, 6rem);
  font-size: 0.72rem;
  font-weight: 750;
  letter-spacing: 0.16em;
  text-transform: uppercase;
}
.publication-brand-mark {
  display: inline-grid;
  place-items: center;
  width: 2rem;
  height: 2rem;
  color: var(--forest);
  background: #d8b46c;
  border-radius: 50%;
  font-size: 0.58rem;
  letter-spacing: 0.03em;
}
.publication-kicker {
  position: relative;
  z-index: 1;
  margin: 0 0 0.7rem;
  color: #e2c27f;
  font-size: 0.72rem;
  font-weight: 800;
  letter-spacing: 0.2em;
  text-transform: uppercase;
}
.publication-hero h1 {
  position: relative;
  z-index: 1;
  max-width: 14ch;
  margin: 0;
  color: inherit;
  font-family: var(--display);
  font-size: clamp(2.7rem, 8vw, 5.4rem);
  font-weight: 500;
  letter-spacing: -0.045em;
  line-height: 0.98;
}
.publication-deck {
  position: relative;
  z-index: 1;
  max-width: 40rem;
  margin: 1.5rem 0 0;
  color: rgb(255 253 248 / 0.78);
  font-family: var(--display);
  font-size: clamp(1.05rem, 2vw, 1.3rem);
  line-height: 1.5;
}
main,
#bibliography {
  width: min(100% - 3rem, 58rem);
  margin-inline: auto;
}
main { padding-block: clamp(2.5rem, 6vw, 5.5rem); }
main > section,
#bibliography {
  margin-block: 1.25rem;
  padding: clamp(1.35rem, 4vw, 2.5rem);
  background: var(--paper-raised);
  border: 1px solid var(--rule);
  border-radius: 1rem;
  box-shadow: 0 0.6rem 1.8rem rgb(40 34 22 / 0.05);
}
main > .publication-meta,
main > .learner-fields,
main > .instructor-banner { margin-top: 0; }
.publication-meta,
.learner-fields {
  display: grid;
  grid-template-columns: repeat(2, minmax(0, 1fr));
  gap: 1rem;
}
.publication-meta p,
.learner-fields p {
  margin: 0;
  padding-bottom: 0.55rem;
  border-bottom: 1px solid var(--rule);
}
h2,
h3,
h4 { text-wrap: balance; }
h2 {
  margin: 0 0 1.15rem;
  color: var(--forest);
  font-family: var(--display);
  font-size: clamp(1.6rem, 4vw, 2.15rem);
  font-weight: 500;
  line-height: 1.15;
}
h2::before {
  content: "";
  display: block;
  width: 2.6rem;
  height: 0.2rem;
  margin-bottom: 0.85rem;
  background: var(--publication-accent);
}
h3 {
  margin: 2rem 0 0.65rem;
  color: var(--ink);
  font-family: var(--display);
  font-size: 1.3rem;
  line-height: 1.25;
}
h4 {
  margin: 1.4rem 0 0.5rem;
  color: var(--publication-accent);
  font-size: 0.72rem;
  letter-spacing: 0.1em;
  text-transform: uppercase;
}
p { margin: 0.7rem 0; }
a {
  color: var(--forest);
  text-decoration-color: #c19b4c;
  text-underline-offset: 0.16em;
}
a:hover { text-decoration-thickness: 0.14em; }
ul,
ol { padding-inline-start: 1.35rem; }
li { margin-block: 0.48rem; padding-inline-start: 0.25rem; }
li::marker { color: var(--publication-accent); font-weight: 750; }
article,
.lesson {
  margin-block: 1.2rem;
  padding: 1.2rem 1.3rem;
  background: #faf7ef;
  border-inline-start: 0.25rem solid var(--publication-accent);
  border-radius: 0.35rem 0.8rem 0.8rem 0.35rem;
}
.section-summary,
.instructor-banner {
  margin-top: 1.25rem;
  padding: 1rem 1.15rem;
  color: #153d2c;
  background: #eef4ee;
  border: 1px solid #cbdcce;
  border-radius: 0.7rem;
}
.instructor-banner {
  color: #673327;
  background: #fbefeb;
  border-color: #e7c8be;
}
.caution-list {
  padding: 1rem 1.2rem 1rem 2.5rem;
  background: #fff8e8;
  border-radius: 0.65rem;
}
pre {
  overflow-x: auto;
  margin: 1.1rem 0;
  padding: 1rem 1.1rem;
  color: #f7f2e8;
  background: #1f2923;
  border-inline-start: 0.25rem solid #d8b46c;
  border-radius: 0.65rem;
}
code { font-family: var(--mono); font-size: 0.9em; }
dl { margin: 0; }
dt { margin-top: 1rem; color: var(--forest); font-weight: 780; }
dd { margin: 0.25rem 0 0; color: var(--muted); }
.definition-grid,
.flashcard-grid {
  display: grid;
  grid-template-columns: repeat(2, minmax(0, 1fr));
  gap: 0.85rem;
}
.definition-grid > div,
.flashcard-grid > div {
  padding: 1rem;
  background: #faf7ef;
  border: 1px solid var(--rule);
  border-radius: 0.65rem;
}
.definition-grid dt,
.flashcard-grid dt { margin-top: 0; }
details {
  margin-top: 0.9rem;
  padding: 0.8rem 1rem;
  background: #f0eee7;
  border-radius: 0.55rem;
}
summary { color: var(--forest); cursor: pointer; font-weight: 760; }
.learner-fields p { min-height: 2.3rem; color: var(--muted); }
.response-space {
  min-height: 7.5rem;
  margin: 1rem 0 0.4rem;
  padding-top: 0.5rem;
  color: var(--muted);
  background: repeating-linear-gradient(
    to bottom,
    transparent 0,
    transparent 1.75rem,
    #ded7c8 1.78rem,
    transparent 1.82rem
  );
  font-size: 0.75rem;
  font-weight: 700;
  letter-spacing: 0.08em;
  text-transform: uppercase;
}
#bibliography { margin-bottom: 4rem; }
.publication-footer {
  padding: 2.25rem 1.5rem;
  color: rgb(255 253 248 / 0.68);
  background: var(--forest);
  text-align: center;
  font-size: 0.72rem;
  letter-spacing: 0.13em;
  text-transform: uppercase;
}
@media (max-width: 42rem) {
  .publication-hero { min-height: 22rem; }
  main,
  #bibliography { width: min(100% - 1.25rem, 58rem); }
  main > section,
  #bibliography { border-radius: 0.7rem; }
  .publication-meta,
  .learner-fields,
  .definition-grid,
  .flashcard-grid { grid-template-columns: 1fr; }
}
@media (prefers-reduced-motion: reduce) {
  html { scroll-behavior: auto; }
}
@media print {
  @page { size: A4; margin: 17mm 16mm 18mm; }
  html,
  body.publication-shell { background: #fff; }
  .publication-hero {
    min-height: 245mm;
    margin: -17mm -16mm 0;
    padding: 25mm 22mm;
    break-after: page;
    print-color-adjust: exact;
  }
  main,
  #bibliography { width: auto; }
  main { padding: 0; }
  main > section,
  #bibliography {
    margin: 0 0 7mm;
    padding: 0;
    border: 0;
    border-radius: 0;
    box-shadow: none;
    break-inside: avoid;
  }
  article,
  .lesson,
  .definition-grid > div,
  .flashcard-grid > div,
  details,
  .section-summary,
  .instructor-banner { break-inside: avoid; print-color-adjust: exact; }
  h2,
  h3,
  h4 { break-after: avoid; }
  a { color: inherit; }
  .publication-footer { display: none; }
}
""".strip()


def publication_theme(publication_kind: PublicationKind) -> PublicationTheme:
    """Return the reviewed theme for one stable publication kind."""

    return _PUBLICATION_THEMES[publication_kind]


def publication_html_start(
    *,
    title: str,
    language: str,
    direction: str,
    publication_kind: PublicationKind,
    escaped_title: str,
) -> str:
    """Build metadata, local CSS, and an accessible editorial cover."""

    theme = publication_theme(publication_kind)
    publication_class = f"publication--{publication_kind}"
    return (
        "<!doctype html>\n"
        f'<html lang="{language}" dir="{direction}">\n'
        "<head>\n"
        '  <meta charset="utf-8">\n'
        '  <meta name="viewport" content="width=device-width,initial-scale=1">\n'
        '  <meta name="color-scheme" content="light">\n'
        f"  <title>{title}</title>\n"
        f'  <style id="txt2crs-publication-theme">{PUBLICATION_CSS}</style>\n'
        "</head>\n"
        f'<body class="publication-shell {publication_class}">\n'
        '<header class="publication-hero">\n'
        '  <div class="publication-brand"><span class="publication-brand-mark" '
        'aria-hidden="true">T2C</span><span>txt2crs / Learning edition</span></div>\n'
        f'  <p class="publication-kicker">{theme.label}</p>\n'
        f"  <h1>{escaped_title}</h1>\n"
        f'  <p class="publication-deck">{theme.deck}</p>\n'
        "</header>\n"
    )


def publication_html_end() -> str:
    """Close a standalone HTML publication with a restrained brand folio."""

    return (
        '<footer class="publication-footer">txt2crs / Built for deliberate learning'
        "</footer>\n</body>\n</html>\n"
    )


def _plain_text_from_markdown(value: str) -> str:
    """Reduce the renderer's trusted inline Markdown subset to plain text."""

    readable_links = _INLINE_MARKDOWN_LINK_PATTERN.sub(r"\1 (\2)", value)
    return readable_links.replace("**", "").replace("__", "").replace("`", "")


def _pdf_lines(value: str, *, width: int) -> list[str]:
    """Wrap one logical line while still breaking unusually long URLs."""

    normalized = _plain_text_from_markdown(value).translate(
        _PDF_PUNCTUATION_TRANSLATION
    )
    return textwrap.wrap(
        normalized,
        width=width,
        break_long_words=True,
        break_on_hyphens=False,
        replace_whitespace=False,
    ) or [""]


def render_publication_pdf(
    *,
    title: str,
    subject: str,
    markdown_lines: list[str],
    publication_kind: PublicationKind,
) -> bytes:
    """Render a searchable A4 publication with a cover and hierarchy."""

    theme = publication_theme(publication_kind)
    document = fitz.open()
    table_of_contents: list[list[int | str]] = [[1, title, 1]]
    page_width = 595.0
    page_height = 842.0
    content_left = 68.0
    content_right = 527.0
    content_top = 88.0
    content_bottom = 785.0

    try:
        # A full-bleed color cover gives every format the same recognizable
        # opening while all text remains searchable and selectable.
        cover_page = document.new_page(width=page_width, height=page_height)
        cover_page.draw_rect(cover_page.rect, color=_FOREST_RGB, fill=_FOREST_RGB)
        cover_page.draw_rect(
            fitz.Rect(0, 0, 16, page_height),
            color=theme.accent_rgb,
            fill=theme.accent_rgb,
        )
        cover_page.draw_circle(
            fitz.Point(86, 92),
            19,
            color=_GOLD_RGB,
            fill=_GOLD_RGB,
        )
        cover_page.insert_text(
            (75, 96),
            "T2C",
            fontname="hebo",
            fontsize=8.5,
            color=_FOREST_RGB,
        )
        cover_page.insert_text(
            (116, 97),
            "TXT2CRS / LEARNING EDITION",
            fontname="hebo",
            fontsize=10.5,
            color=(1.0, 0.99, 0.965),
        )
        cover_page.insert_text(
            (68, 255),
            theme.label.upper(),
            fontname="hebo",
            fontsize=10,
            color=_GOLD_RGB,
        )
        cover_page.insert_textbox(
            fitz.Rect(68, 280, 525, 545),
            _plain_text_from_markdown(title).translate(_PDF_PUNCTUATION_TRANSLATION),
            fontname="hebo",
            fontsize=31,
            lineheight=1.04,
            color=(1.0, 0.99, 0.965),
        )
        cover_page.insert_textbox(
            fitz.Rect(68, 615, 475, 700),
            theme.deck,
            fontname="heit",
            fontsize=13,
            lineheight=1.35,
            color=(0.865, 0.835, 0.765),
        )
        cover_page.draw_line(
            fitz.Point(68, 760),
            fitz.Point(190, 760),
            color=theme.accent_rgb,
            width=2,
        )
        cover_page.insert_text(
            (68, 785),
            "RESEARCH-GROUNDED / DETERMINISTIC / READY TO TEACH",
            fontname="hebo",
            fontsize=7.5,
            color=(0.78, 0.76, 0.70),
        )

        current_page = document.new_page(width=page_width, height=page_height)
        vertical_position = content_top

        def prepare_content_page(page: fitz.Page) -> None:
            """Draw stable furniture on each reading page."""

            page.draw_rect(page.rect, color=(1.0, 1.0, 1.0), fill=(1.0, 1.0, 1.0))
            page.draw_rect(
                fitz.Rect(0, 0, 6, page_height),
                color=theme.accent_rgb,
                fill=theme.accent_rgb,
            )
            page.insert_text(
                (content_left, 45),
                "txt2crs",
                fontname="hebo",
                fontsize=9,
                color=_FOREST_RGB,
            )
            page.insert_textbox(
                fitz.Rect(300, 33, content_right, 49),
                theme.label.upper(),
                fontname="hebo",
                fontsize=7.5,
                color=theme.accent_rgb,
                align=fitz.TEXT_ALIGN_RIGHT,
            )
            page.draw_line(
                fitz.Point(content_left, 58),
                fitz.Point(content_right, 58),
                color=_RULE_RGB,
                width=0.6,
            )

        prepare_content_page(current_page)

        def begin_new_page() -> None:
            """Advance the local layout cursor to a fresh content page."""

            nonlocal current_page, vertical_position
            current_page = document.new_page(width=page_width, height=page_height)
            prepare_content_page(current_page)
            vertical_position = content_top

        def require_space(height: float) -> None:
            """Keep a heading or short block above the footer boundary."""

            if vertical_position + height > content_bottom:
                begin_new_page()

        def draw_text_block(
            text: str,
            *,
            font_name: str,
            font_size: float,
            color: tuple[float, float, float],
            leading: float,
            wrap_width: int,
            left_indent: float = 0,
            prefix: str = "",
            background: tuple[float, float, float] | None = None,
        ) -> None:
            """Draw wrapped searchable lines and paginate them as needed."""

            nonlocal vertical_position
            wrapped_lines = _pdf_lines(text, width=wrap_width)
            for line_index, wrapped_line in enumerate(wrapped_lines):
                require_space(leading + 2)
                rendered_prefix = prefix if line_index == 0 else ""
                if background is not None:
                    current_page.draw_rect(
                        fitz.Rect(
                            content_left + left_indent - 7,
                            vertical_position - font_size,
                            content_right,
                            vertical_position + 5,
                        ),
                        color=background,
                        fill=background,
                    )
                current_page.insert_text(
                    (content_left + left_indent, vertical_position),
                    f"{rendered_prefix}{wrapped_line}",
                    fontname=font_name,
                    fontsize=font_size,
                    color=color,
                )
                vertical_position += leading

        def draw_assessment_response_space() -> None:
            """Add clearly ruled writing space after one PDF assessment item."""

            nonlocal vertical_position
            require_space(118)
            vertical_position += 8
            current_page.insert_text(
                (content_left, vertical_position),
                "RESPONSE NOTES",
                fontname="hebo",
                fontsize=7.5,
                color=theme.accent_rgb,
            )
            vertical_position += 13
            for _ in range(5):
                current_page.draw_line(
                    fitz.Point(content_left, vertical_position),
                    fitz.Point(content_right, vertical_position),
                    color=_RULE_RGB,
                    width=0.55,
                )
                vertical_position += 18

        if publication_kind == "assessment":
            # Paper assessments need a practical identity strip before the
            # instructions. The lines remain vector shapes and print cleanly.
            current_page.insert_text(
                (content_left, vertical_position),
                "STUDENT NAME",
                fontname="hebo",
                fontsize=7.5,
                color=_MUTED_RGB,
            )
            current_page.insert_text(
                (338, vertical_position),
                "DATE",
                fontname="hebo",
                fontsize=7.5,
                color=_MUTED_RGB,
            )
            vertical_position += 12
            current_page.draw_line(
                fitz.Point(content_left, vertical_position),
                fitz.Point(315, vertical_position),
                color=_RULE_RGB,
                width=0.7,
            )
            current_page.draw_line(
                fitz.Point(338, vertical_position),
                fitz.Point(content_right, vertical_position),
                color=_RULE_RGB,
                width=0.7,
            )
            vertical_position += 27

        inside_code_block = False
        skipped_cover_heading = False
        assessment_question_open = False
        for original_line in markdown_lines:
            stripped_line = original_line.strip()
            if stripped_line == "```":
                inside_code_block = not inside_code_block
                vertical_position += 4
                continue
            if inside_code_block:
                draw_text_block(
                    original_line,
                    font_name="cour",
                    font_size=8.7,
                    color=(0.91, 0.93, 0.91),
                    leading=13,
                    wrap_width=76,
                    left_indent=10,
                    background=(0.12, 0.16, 0.14),
                )
                continue
            if not stripped_line:
                vertical_position += 5
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped_line)
            if heading_match is not None:
                heading_level = len(heading_match.group(1))
                heading_text = _plain_text_from_markdown(heading_match.group(2))
                if heading_level == 1 and not skipped_cover_heading:
                    skipped_cover_heading = True
                    continue
                if publication_kind == "assessment" and heading_level == 2:
                    if assessment_question_open:
                        draw_assessment_response_space()
                    assessment_question_open = True
                heading_size = {1: 22.0, 2: 17.5, 3: 12.5}.get(
                    heading_level,
                    10.0,
                )
                heading_leading = heading_size * 1.25
                heading_color = _FOREST_RGB if heading_level <= 2 else _INK_RGB
                vertical_position += 13 if heading_level <= 2 else 8
                require_space(heading_leading * 2)
                table_of_contents.append(
                    [min(heading_level, 3), heading_text, current_page.number + 1]
                )
                if heading_level <= 2:
                    current_page.draw_rect(
                        fitz.Rect(
                            content_left,
                            vertical_position - heading_size - 5,
                            content_left + 34,
                            vertical_position - heading_size - 2,
                        ),
                        color=theme.accent_rgb,
                        fill=theme.accent_rgb,
                    )
                draw_text_block(
                    heading_text,
                    font_name="hebo",
                    font_size=heading_size,
                    color=heading_color,
                    leading=heading_leading,
                    wrap_width=54 if heading_level <= 2 else 72,
                )
                vertical_position += 4
                continue

            bullet_match = re.match(r"^[-*]\s+(.+)$", stripped_line)
            numbered_match = re.match(r"^(\d+)\.\s+(.+)$", stripped_line)
            option_match = re.match(r"^([A-Z])\.\s+(.+)$", stripped_line)
            if bullet_match is not None:
                draw_text_block(
                    bullet_match.group(1),
                    font_name="helv",
                    font_size=10.2,
                    color=_INK_RGB,
                    leading=15,
                    wrap_width=78,
                    left_indent=18,
                    prefix="-  ",
                )
                vertical_position += 2
                continue
            if numbered_match is not None or option_match is not None:
                list_match = numbered_match or option_match
                assert list_match is not None
                draw_text_block(
                    list_match.group(2),
                    font_name="helv",
                    font_size=10.2,
                    color=_INK_RGB,
                    leading=15,
                    wrap_width=77,
                    left_indent=18,
                    prefix=f"{list_match.group(1)}.  ",
                )
                vertical_position += 2
                continue

            is_callout = stripped_line.startswith(
                ("**Summary:**", "**Answer:**", "**Solution:**", "**Worked solution:**")
            )
            draw_text_block(
                stripped_line,
                font_name="hebo" if is_callout else "helv",
                font_size=10.2,
                color=_FOREST_RGB if is_callout else _INK_RGB,
                leading=15.5,
                wrap_width=84 if is_callout else 88,
                left_indent=8 if is_callout else 0,
                background=(0.925, 0.957, 0.925) if is_callout else None,
            )
            vertical_position += 3

        if publication_kind == "assessment" and assessment_question_open:
            draw_assessment_response_space()

        # Add folios after pagination so each page has an accurate total.
        total_pages = len(document)
        for page_index in range(1, total_pages):
            page = document[page_index]
            page.draw_line(
                fitz.Point(content_left, 806),
                fitz.Point(content_right, 806),
                color=_RULE_RGB,
                width=0.5,
            )
            page.insert_text(
                (content_left, 824),
                "txt2crs / learning edition",
                fontname="helv",
                fontsize=8.2,
                color=_MUTED_RGB,
            )
            page.insert_textbox(
                fitz.Rect(420, 812, content_right, 828),
                f"{page_index + 1} / {total_pages}",
                fontname="hebo",
                fontsize=8.2,
                color=theme.accent_rgb,
                align=fitz.TEXT_ALIGN_RIGHT,
            )

        # Preserve native click targets wherever a complete printed URL fits
        # on one line. Long wrapped URLs remain readable text even when a PDF
        # viewer cannot expose them as one annotation rectangle.
        source_urls = {
            link_match.group(2)
            for markdown_line in markdown_lines
            for link_match in _INLINE_MARKDOWN_LINK_PATTERN.finditer(markdown_line)
        }
        for page_index in range(1, total_pages):
            page = document[page_index]
            for source_url in sorted(source_urls):
                for source_rectangle in page.search_for(source_url):
                    page.insert_link(
                        {
                            "kind": fitz.LINK_URI,
                            "from": source_rectangle,
                            "uri": source_url,
                        }
                    )

        document.set_toc(table_of_contents)
        document.set_metadata(
            {
                "title": title,
                "subject": subject,
                "author": "txt2crs",
                "creator": "txt2crs deterministic publication renderer",
                "keywords": f"education, {theme.label.casefold()}, txt2crs",
            }
        )
        return bytes(document.tobytes(garbage=4, deflate=True))
    finally:
        document.close()


def _set_cell_shading(cell: object, fill_hex: str) -> None:
    """Apply a solid OOXML fill to a table cell used as a design panel."""

    cell_properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell_properties.append(shading)


def _set_cell_margins(
    cell: object,
    *,
    top: int,
    start: int,
    bottom: int,
    end: int,
) -> None:
    """Set explicit cell padding because python-docx has no public shortcut."""

    cell_properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        margin_element = OxmlElement(f"w:{margin_name}")
        margin_element.set(qn("w:w"), str(margin_value))
        margin_element.set(qn("w:type"), "dxa")
        margins.append(margin_element)


def _set_run_font(run: object, font_name: str) -> None:
    """Set every OOXML font slot so Word fallback stays predictable."""

    run.font.name = font_name  # type: ignore[attr-defined]
    run_properties = run._element.get_or_add_rPr()  # type: ignore[attr-defined]
    run_fonts = run_properties.rFonts
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)
    for font_attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        run_fonts.set(qn(f"w:{font_attribute}"), font_name)


def _set_style_font(style: object, font_name: str) -> None:
    """Apply a font through python-docx and the raw OOXML style layer."""

    style.font.name = font_name  # type: ignore[attr-defined]
    style_properties = style._element.get_or_add_rPr()  # type: ignore[attr-defined]
    run_fonts = style_properties.rFonts
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        style_properties.insert(0, run_fonts)
    for font_attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        run_fonts.set(qn(f"w:{font_attribute}"), font_name)


def _add_field(paragraph: Paragraph, instruction: str, placeholder: str) -> None:
    """Append a Word field that updates when the file opens or prints."""

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction_text, separate, value, end])


def _configure_docx_styles(document: DocxDocument, theme: PublicationTheme) -> None:
    """Define the native Word style vocabulary before adding content."""

    styles = document.styles
    normal_style = styles["Normal"]
    _set_style_font(normal_style, "Aptos")
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x21, 0x1F, 0x1B)
    normal_style.paragraph_format.space_after = Pt(7)
    normal_style.paragraph_format.line_spacing = 1.22
    normal_style.paragraph_format.widow_control = True

    title_style = styles["Title"]
    _set_style_font(title_style, "Georgia")
    title_style.font.size = Pt(31)
    title_style.font.bold = False
    title_style.font.color.rgb = RGBColor(0x1A, 0x50, 0x38)

    heading_specs = {
        "Heading 1": ("Georgia", 22, _FOREST_HEX, 18, 8),
        "Heading 2": ("Georgia", 16, theme.accent_hex, 14, 6),
        "Heading 3": ("Aptos Display", 12, "292720", 11, 4),
        "Heading 4": ("Aptos", 9, theme.accent_hex, 9, 3),
    }
    for style_name, (
        font_name,
        font_size,
        color_hex,
        before_spacing,
        after_spacing,
    ) in heading_specs.items():
        heading_style = styles[style_name]
        _set_style_font(heading_style, font_name)
        heading_style.font.size = Pt(font_size)
        heading_style.font.bold = style_name != "Heading 1"
        heading_style.font.color.rgb = RGBColor.from_string(color_hex)
        heading_style.paragraph_format.space_before = Pt(before_spacing)
        heading_style.paragraph_format.space_after = Pt(after_spacing)
        heading_style.paragraph_format.keep_with_next = True
        heading_style.paragraph_format.keep_together = True
        heading_style.paragraph_format.widow_control = True

    custom_style_specs = {
        "Publication Meta": ("Aptos", 8.5, "6A655A"),
        "Publication Callout": ("Aptos", 10.5, _FOREST_HEX),
        "Publication Code": ("Consolas", 9, "F5F1E7"),
        "Publication Contents": ("Aptos", 10, "38352E"),
        "Publication Response": ("Aptos", 9, "8B857A"),
    }
    for style_name, (
        custom_font_name,
        custom_font_size,
        custom_color_hex,
    ) in custom_style_specs.items():
        style = (
            styles[style_name]
            if style_name in styles
            else styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        )
        _set_style_font(style, custom_font_name)
        style.font.size = Pt(custom_font_size)
        style.font.color.rgb = RGBColor.from_string(custom_color_hex)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.widow_control = True

    callout_style = styles["Publication Callout"]
    callout_style.paragraph_format.left_indent = Inches(0.16)
    callout_style.paragraph_format.right_indent = Inches(0.12)
    callout_style.paragraph_format.space_before = Pt(5)
    callout_style.paragraph_format.space_after = Pt(8)
    callout_properties = callout_style._element.get_or_add_pPr()
    callout_shading = OxmlElement("w:shd")
    callout_shading.set(qn("w:fill"), "EEF4EE")
    callout_properties.append(callout_shading)

    code_style = styles["Publication Code"]
    code_style.paragraph_format.left_indent = Inches(0.16)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_after = Pt(0)
    code_properties = code_style._element.get_or_add_pPr()
    code_shading = OxmlElement("w:shd")
    code_shading.set(qn("w:fill"), "243029")
    code_properties.append(code_shading)


def _configure_docx_page(document: DocxDocument, theme: PublicationTheme) -> None:
    """Set print geometry, running furniture, and page fields."""

    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.text = f"txt2crs  /  {theme.label}"
    header_paragraph.style = document.styles["Publication Meta"]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_border = OxmlElement("w:pBdr")
    bottom_border = OxmlElement("w:bottom")
    bottom_border.set(qn("w:val"), "single")
    bottom_border.set(qn("w:sz"), "6")
    bottom_border.set(qn("w:space"), "5")
    bottom_border.set(qn("w:color"), theme.accent_hex)
    header_border.append(bottom_border)
    header_paragraph._p.get_or_add_pPr().append(header_border)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.style = document.styles["Publication Meta"]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.add_run("txt2crs  /  ")
    _add_field(footer_paragraph, "PAGE", "1")
    footer_paragraph.add_run("  of  ")
    _add_field(footer_paragraph, "NUMPAGES", "1")

    # Ask compatible editors to refresh total-page fields on open.
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _add_docx_cover(
    document: DocxDocument,
    *,
    title: str,
    theme: PublicationTheme,
) -> None:
    """Create a color-safe cover panel from native Word table cells."""

    cover_table = document.add_table(rows=1, cols=1)
    cover_table.autofit = False
    cover_table.columns[0].width = Inches(6.75)
    cover_row = cover_table.rows[0]
    cover_row.height = Inches(8.15)
    cover_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cover_cell = cover_row.cells[0]
    cover_cell.width = Inches(6.75)
    cover_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_shading(cover_cell, _FOREST_HEX)
    _set_cell_margins(cover_cell, top=380, start=430, bottom=380, end=430)

    brand_paragraph = cover_cell.paragraphs[0]
    brand_paragraph.paragraph_format.space_after = Pt(76)
    brand_run = brand_paragraph.add_run("TXT2CRS  /  LEARNING EDITION")
    _set_run_font(brand_run, "Aptos")
    brand_run.font.size = Pt(9)
    brand_run.font.bold = True
    brand_run.font.color.rgb = RGBColor(0xE2, 0xC2, 0x7F)

    kicker_paragraph = cover_cell.add_paragraph()
    kicker_paragraph.paragraph_format.space_after = Pt(8)
    kicker_run = kicker_paragraph.add_run(theme.label.upper())
    _set_run_font(kicker_run, "Aptos")
    kicker_run.font.size = Pt(9)
    kicker_run.font.bold = True
    kicker_run.font.color.rgb = RGBColor(0xE2, 0xC2, 0x7F)

    title_paragraph = cover_cell.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(18)
    title_run = title_paragraph.add_run(_plain_text_from_markdown(title))
    _set_run_font(title_run, "Georgia")
    title_run.font.size = Pt(29)
    title_run.font.color.rgb = RGBColor(0xFF, 0xFD, 0xF8)

    deck_paragraph = cover_cell.add_paragraph()
    deck_run = deck_paragraph.add_run(theme.deck)
    _set_run_font(deck_run, "Georgia")
    deck_run.font.size = Pt(11)
    deck_run.font.italic = True
    deck_run.font.color.rgb = RGBColor(0xD9, 0xD1, 0xC1)

    cover_note = document.add_paragraph(style="Publication Meta")
    cover_note.paragraph_format.space_before = Pt(16)
    cover_note.paragraph_format.space_after = Pt(0)
    cover_note.add_run("RESEARCH-GROUNDED  /  DETERMINISTIC  /  READY TO TEACH")

    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)


def _markdown_headings(markdown_content: str) -> list[tuple[int, str]]:
    """Extract visible heading structure for a compact contents page."""

    headings: list[tuple[int, str]] = []
    for line in markdown_content.splitlines():
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line.strip())
        if heading_match is not None:
            headings.append(
                (
                    len(heading_match.group(1)),
                    _plain_text_from_markdown(heading_match.group(2)),
                )
            )
    return headings


def _add_docx_contents(document: DocxDocument, markdown_content: str) -> None:
    """Add an immediately readable map before the document body."""

    headings = _markdown_headings(markdown_content)
    # A static contents page is valuable for long courses and review packs but
    # feels wasteful for a short worksheet or a one-item answer key. Native
    # Word headings still provide navigation in those compact publications.
    if len(headings) < 3:
        return

    document.add_heading("Inside this publication", level=1)
    visible_headings = headings[:16]
    for heading_index, (heading_level, heading_text) in enumerate(
        visible_headings,
        start=1,
    ):
        paragraph = document.add_paragraph(style="Publication Contents")
        paragraph.paragraph_format.left_indent = Inches(
            0.18 * max(0, heading_level - 2)
        )
        number_run = paragraph.add_run(f"{heading_index:02d}  ")
        number_run.font.bold = True
        number_run.font.color.rgb = RGBColor(0xA0, 0x6D, 0x18)
        paragraph.add_run(heading_text)
    if len(headings) > len(visible_headings):
        remaining_paragraph = document.add_paragraph(style="Publication Meta")
        remaining_paragraph.add_run(
            f"Plus {len(headings) - len(visible_headings)} additional sections."
        )
    contents_break = document.add_paragraph()
    contents_break.add_run().add_break(WD_BREAK.PAGE)


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    """Add a native external hyperlink instead of inert link syntax."""

    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _FOREST_HEX)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    link_text = OxmlElement("w:t")
    link_text.text = text
    run.append(link_text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline_docx_content(paragraph: Paragraph, value: str) -> None:
    """Render links, emphasis, and inline code as native Word runs."""

    cursor = 0
    for token_match in _INLINE_DOCX_TOKEN_PATTERN.finditer(value):
        if token_match.start() > cursor:
            paragraph.add_run(value[cursor : token_match.start()])
        token = token_match.group(0)
        link_match = _INLINE_MARKDOWN_LINK_PATTERN.fullmatch(token)
        if link_match is not None:
            _add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
            # Keep the target visible for printed copies and plain-text
            # extraction while also preserving the native clickable relation.
            paragraph.add_run(f" ({link_match.group(2)})")
        elif token.startswith("**"):
            emphasized_run = paragraph.add_run(token[2:-2])
            emphasized_run.font.bold = True
        else:
            code_run = paragraph.add_run(token[1:-1])
            _set_run_font(code_run, "Consolas")
            code_run.font.size = Pt(9)
        cursor = token_match.end()
    if cursor < len(value):
        paragraph.add_run(value[cursor:])


def _add_assessment_fields(document: DocxDocument) -> None:
    """Create learner fields at the start of a printable assessment."""

    fields_table = document.add_table(rows=1, cols=2)
    fields_table.autofit = False
    labels = ("Student name", "Date")
    for field_cell, field_label in zip(fields_table.rows[0].cells, labels, strict=True):
        _set_cell_shading(field_cell, "F5F1E7")
        _set_cell_margins(field_cell, top=120, start=140, bottom=240, end=140)
        field_paragraph = field_cell.paragraphs[0]
        field_paragraph.style = document.styles["Publication Meta"]
        field_paragraph.add_run(f"{field_label}\n\n")
        field_paragraph.add_run("________________________________")
    document.add_paragraph()


def _add_assessment_response_space(document: DocxDocument) -> None:
    """Append a bounded writing area after one assessment item."""

    response_label = document.add_paragraph("Response notes", style="Publication Meta")
    response_label.paragraph_format.space_before = Pt(9)
    for _ in range(3):
        response_line = document.add_paragraph(style="Publication Response")
        response_line.paragraph_format.space_after = Pt(6)
        response_line.add_run(
            "____________________________________________________________"
        )


def render_publication_docx(
    *,
    title: str,
    subject: str,
    markdown_content: str,
    publication_kind: PublicationKind,
) -> bytes:
    """Create a styled, editable, and navigable Word publication."""

    theme = publication_theme(publication_kind)
    document = Document()
    document.core_properties.title = title
    document.core_properties.subject = subject
    document.core_properties.author = "txt2crs"
    document.core_properties.category = "Education"
    document.core_properties.keywords = f"txt2crs, {theme.label}, education"
    _configure_docx_styles(document, theme)
    _configure_docx_page(document, theme)
    _add_docx_cover(document, title=title, theme=theme)
    _add_docx_contents(document, markdown_content)
    if publication_kind == "assessment":
        _add_assessment_fields(document)
    if publication_kind == "answer-key":
        instructor_notice = document.add_paragraph(style="Publication Callout")
        instructor_notice.add_run(
            "Instructor edition: keep this publication separate from learner materials."
        )

    inside_code_block = False
    skipped_cover_heading = False
    assessment_question_open = False
    for raw_line in markdown_content.splitlines():
        stripped_line = raw_line.strip()
        if stripped_line == "```":
            inside_code_block = not inside_code_block
            continue
        if inside_code_block:
            code_paragraph = document.add_paragraph(style="Publication Code")
            code_paragraph.add_run(raw_line or " ")
            continue
        if not stripped_line:
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped_line)
        if heading_match is not None:
            heading_level = len(heading_match.group(1))
            if heading_level == 1 and not skipped_cover_heading:
                skipped_cover_heading = True
                continue
            if publication_kind == "assessment" and heading_level == 2:
                if assessment_question_open:
                    _add_assessment_response_space(document)
                assessment_question_open = True
            heading = document.add_heading(
                _plain_text_from_markdown(heading_match.group(2)),
                level=min(heading_level, 4),
            )
            heading.paragraph_format.keep_with_next = True
            heading.paragraph_format.keep_together = True
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped_line)
        numbered_match = re.match(r"^\d+\.\s+(.+)$", stripped_line)
        option_match = re.match(r"^([A-Z])\.\s+(.+)$", stripped_line)
        if bullet_match is not None:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_docx_content(paragraph, bullet_match.group(1))
            continue
        if numbered_match is not None:
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_docx_content(paragraph, numbered_match.group(1))
            continue
        if option_match is not None:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            option_label = paragraph.add_run(f"{option_match.group(1)}.  ")
            option_label.font.bold = True
            _add_inline_docx_content(paragraph, option_match.group(2))
            continue

        is_callout = stripped_line.startswith(
            ("**Summary:**", "**Answer:**", "**Solution:**", "**Worked solution:**")
        )
        paragraph = document.add_paragraph(
            style="Publication Callout" if is_callout else None
        )
        _add_inline_docx_content(paragraph, stripped_line)

    if publication_kind == "assessment" and assessment_question_open:
        _add_assessment_response_space(document)

    output_buffer = BytesIO()
    document.save(output_buffer)
    return output_buffer.getvalue()
