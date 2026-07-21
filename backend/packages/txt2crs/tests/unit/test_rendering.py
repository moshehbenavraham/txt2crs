# SPDX-License-Identifier: MIT-0

"""Tests for deterministic, accessible, active-content-free rendering."""

from io import BytesIO

import fitz  # type: ignore[import-untyped]
import pytest
from docx import Document
from docx.shared import Inches, RGBColor

from tests.factories import (
    valid_answer_key_data,
    valid_assessment_data,
    valid_course_data,
    valid_review_pack_data,
)
from txt2crs.domain.models import AnswerKey, Assessment, Course, ReviewPack
from txt2crs.domain.validation import ArtifactBundle, validate_artifact_bundle
from txt2crs.rendering.artifacts import (
    ArtifactRenderer,
    RenderedOutputQaError,
    derive_safe_filename,
    validate_rendered_html,
)


def valid_bundle() -> ArtifactBundle:
    """Return the cross-validated artifact bundle used by render tests."""

    return validate_artifact_bundle(
        course=Course.model_validate(valid_course_data()),
        review_pack=ReviewPack.model_validate(valid_review_pack_data()),
        assessment=Assessment.model_validate(valid_assessment_data()),
        answer_key=AnswerKey.model_validate(valid_answer_key_data()),
    )


def test_filename_is_deterministic_ascii_and_path_safe() -> None:
    """Filename generation never needs a model and cannot traverse paths."""

    assert derive_safe_filename("../../Python: Variables?!") == "python-variables"
    assert derive_safe_filename("  \u05e7\u05d5\u05e8\u05e1 Python  ") == "python"
    assert derive_safe_filename("...") == "course"


def test_course_html_escapes_untrusted_content_and_is_semantic() -> None:
    """Model text is data: scripts render as text, not active markup."""

    bundle = valid_bundle()
    bundle.course.modules[0].sections[0].content_blocks[
        0
    ].text = "<script>alert('unsafe')</script> A variable."
    renderer = ArtifactRenderer()

    rendered = renderer.render_bundle(bundle)
    course_html = rendered["course_html"].content.decode("utf-8")

    assert '<html lang="en" dir="ltr">' in course_html
    assert '<main id="course-content">' in course_html
    assert "<h1>Python Basics</h1>" in course_html
    assert "&lt;script&gt;" in course_html
    assert "<script>" not in course_html
    assert 'id="bibliography"' in course_html
    validate_rendered_html(
        course_html,
        required_element_ids={"course-content", "bibliography"},
    )


def test_html_publications_embed_the_responsive_print_design_system() -> None:
    """Every standalone HTML file should feel intentionally publication-ready."""

    rendered = ArtifactRenderer().render_bundle(valid_bundle())

    expected_publication_classes = {
        "course_html": "publication--course",
        "review_pack_html": "publication--review",
        "assessment_html": "publication--assessment",
        "answer_key_html": "publication--answer-key",
    }
    for artifact_name, publication_class in expected_publication_classes.items():
        html_content = rendered[artifact_name].content.decode("utf-8")

        assert '<style id="txt2crs-publication-theme">' in html_content
        assert f'class="publication-shell {publication_class}"' in html_content
        assert 'class="publication-hero"' in html_content
        assert 'class="publication-brand"' in html_content
        assert 'class="publication-kicker"' in html_content
        assert "@media print" in html_content
        assert "@media (max-width: 42rem)" in html_content
        assert "--publication-accent" in html_content
        assert "prefers-reduced-motion" in html_content
        assert "overflow-wrap: anywhere" in html_content


def test_assessment_includes_printable_learner_fields_and_response_space() -> None:
    """The assessment should function as a polished screen and print worksheet."""

    rendered = ArtifactRenderer().render_bundle(valid_bundle())
    assessment_html = rendered["assessment_html"].content.decode("utf-8")
    assessment_pdf = fitz.open(
        stream=rendered["assessment_pdf"].content,
        filetype="pdf",
    )
    try:
        assessment_pdf_text = "\n".join(
            page.get_text("text") for page in assessment_pdf
        )
    finally:
        assessment_pdf.close()
    assessment_document = Document(BytesIO(rendered["assessment_docx"].content))
    assessment_docx_text = "\n".join(
        [paragraph.text for paragraph in assessment_document.paragraphs]
        + [
            paragraph.text
            for table in assessment_document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        ]
    )

    assert 'class="learner-fields"' in assessment_html
    assert 'class="response-space"' in assessment_html
    assert "Student name" in assessment_html
    assert "STUDENT NAME" in assessment_pdf_text
    assert "RESPONSE NOTES" in assessment_pdf_text
    assert "Student name" in assessment_docx_text
    assert "Response notes" in assessment_docx_text


def test_student_assessment_never_contains_instructor_answers() -> None:
    """Answers, explanations, and rubrics stay in the instructor artifact."""

    rendered = ArtifactRenderer().render_bundle(valid_bundle())
    student_html = rendered["assessment_html"].content.decode("utf-8")
    instructor_html = rendered["answer_key_html"].content.decode("utf-8")

    assert "class_size = 24" not in student_html
    assert "class_size = 24" in instructor_html
    assert "Correct assignment" in instructor_html


def test_answer_key_discloses_evidence_sources_without_leaking_them_to_student() -> (
    None
):
    """Instructor answers retain usable evidence links from canonical course data."""

    rendered = ArtifactRenderer().render_bundle(valid_bundle())
    student_html = rendered["assessment_html"].content.decode("utf-8")
    instructor_html = rendered["answer_key_html"].content.decode("utf-8")
    instructor_markdown = rendered["answer_key_markdown"].content.decode("utf-8")

    assert "Evidence sources" not in student_html
    assert "The Python Tutorial" not in student_html
    for instructor_text in (instructor_html, instructor_markdown):
        assert "Evidence sources" in instructor_text
        assert "The Python Tutorial" in instructor_text
        assert "https://docs.python.org/3/tutorial/" in instructor_text


def test_every_deliverable_has_html_markdown_searchable_pdf_and_docx() -> None:
    """Learners and instructors receive portable formats for every artifact."""

    rendered = ArtifactRenderer().render_bundle(valid_bundle())

    assert set(rendered) == {
        "course_html",
        "course_markdown",
        "course_pdf",
        "course_docx",
        "review_pack_html",
        "review_pack_markdown",
        "review_pack_pdf",
        "review_pack_docx",
        "assessment_html",
        "assessment_markdown",
        "assessment_pdf",
        "assessment_docx",
        "answer_key_html",
        "answer_key_markdown",
        "answer_key_pdf",
        "answer_key_docx",
    }
    for artifact_name in (
        "course_pdf",
        "review_pack_pdf",
        "assessment_pdf",
        "answer_key_pdf",
    ):
        pdf_document = fitz.open(
            stream=rendered[artifact_name].content,
            filetype="pdf",
        )
        try:
            extracted_text = "\n".join(page.get_text("text") for page in pdf_document)
        finally:
            pdf_document.close()
        assert extracted_text.strip()

    for artifact_name in (
        "course_docx",
        "review_pack_docx",
        "assessment_docx",
        "answer_key_docx",
    ):
        document = Document(BytesIO(rendered[artifact_name].content))
        extracted_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert extracted_text.strip()


def test_pdf_publications_have_cover_hierarchy_outlines_and_folios() -> None:
    """PDFs should be designed publications, not fixed-width text dumps."""

    rendered = ArtifactRenderer().render_bundle(valid_bundle())

    for artifact_name in (
        "course_pdf",
        "review_pack_pdf",
        "assessment_pdf",
        "answer_key_pdf",
    ):
        pdf_document = fitz.open(
            stream=rendered[artifact_name].content,
            filetype="pdf",
        )
        try:
            assert len(pdf_document) >= 2
            assert "TXT2CRS" in pdf_document[0].get_text("text").upper()
            assert any(
                drawing.get("fill") is not None
                for drawing in pdf_document[0].get_drawings()
            )
            assert pdf_document.get_toc()

            content_spans = [
                span
                for page in list(pdf_document)[1:]
                for block in page.get_text("dict")["blocks"]
                if "lines" in block
                for line in block["lines"]
                for span in line["spans"]
            ]
            assert max(float(span["size"]) for span in content_spans) >= 16
            assert min(float(span["size"]) for span in content_spans) <= 10.5

            total_pages = len(pdf_document)
            for page_number, page in enumerate(list(pdf_document)[1:], start=2):
                page_text = page.get_text("text")
                assert "txt2crs" in page_text.casefold()
                assert f"{page_number} / {total_pages}" in page_text

            if artifact_name in {
                "course_pdf",
                "review_pack_pdf",
                "answer_key_pdf",
            }:
                link_targets = {
                    link["uri"]
                    for page in pdf_document
                    for link in page.get_links()
                    if link.get("uri") is not None
                }
                assert "https://docs.python.org/3/tutorial/" in link_targets
        finally:
            pdf_document.close()


def test_docx_publications_use_a_branded_navigable_template() -> None:
    """DOCX files should retain premium styling and useful Word semantics."""

    rendered = ArtifactRenderer().render_bundle(valid_bundle())

    for artifact_name in (
        "course_docx",
        "review_pack_docx",
        "assessment_docx",
        "answer_key_docx",
    ):
        document = Document(BytesIO(rendered[artifact_name].content))
        document_xml = document._element.xml
        relationship_targets = {
            relationship.target_ref
            for relationship in document.part.rels.values()
            if relationship.is_external
        }

        assert document.styles["Title"].font.color.rgb == RGBColor(0x1A, 0x50, 0x38)
        assert "Publication Meta" in document.styles
        assert "Publication Callout" in document.styles
        assert "Publication Code" in document.styles
        assert document.sections[0].top_margin == Inches(0.7)
        page_height = document.sections[0].page_height
        assert page_height is not None
        assert abs(page_height - Inches(11.69)) < Inches(0.01)
        assert document.sections[0].different_first_page_header_footer
        assert "txt2crs" in document.sections[0].header.paragraphs[0].text.casefold()
        assert "PAGE" in document.sections[0].footer._element.xml
        assert "NUMPAGES" in document.sections[0].footer._element.xml
        document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        if artifact_name in {"course_docx", "review_pack_docx"}:
            assert "Inside this publication" in document_text
        else:
            assert "Inside this publication" not in document_text
        assert "1A5038" in document_xml
        assert 'w:type="page"' in document_xml
        assert document.tables[0].rows[0].height is not None
        assert document.tables[0].rows[0].height > Inches(7.5)
        assert any(
            paragraph.style is not None
            and paragraph.style.name.startswith("Heading")
            and paragraph.paragraph_format.keep_with_next
            for paragraph in document.paragraphs
        )

        # Only publications containing source links need external hyperlink
        # relationships, but at least those files must preserve live targets.
        if artifact_name in {"course_docx", "review_pack_docx", "answer_key_docx"}:
            assert "https://docs.python.org/3/tutorial/" in relationship_targets


def test_renderers_include_every_canonical_course_and_review_section() -> None:
    """A validated field cannot disappear merely because rendering succeeded."""

    rendered = ArtifactRenderer().render_bundle(valid_bundle())
    course_html = rendered["course_html"].content.decode("utf-8")
    review_html = rendered["review_pack_html"].content.decode("utf-8")

    for expected_course_text in (
        "Basic computer literacy",
        "This module introduces Python variables.",
        "A variable is not a permanent storage location.",
        "student_count = 24",
        "A name bound to a value.",
    ):
        assert expected_course_text in course_html
    for expected_review_text in (
        "The equals sign is assignment",
        "A name bound to a value.",
        "Store the class size.",
        "Variables name values.",
        "Read the guide",
        "The Python Tutorial",
    ):
        assert expected_review_text in review_html


def test_review_renderers_replace_internal_identifiers_with_reader_labels() -> None:
    """Canonical IDs support validation but must not become learner-facing prose."""

    bundle = valid_bundle()
    bundle.review_pack.review_sequence = [
        (
            "Review `sec-variables`, use the `obj-variables` flashcards, "
            "check `card-variable`, then complete `worked-variable` and "
            "`practice-variable`."
        )
    ]

    rendered = ArtifactRenderer().render_bundle(bundle)
    review_html = rendered["review_pack_html"].content.decode("utf-8")
    review_markdown = rendered["review_pack_markdown"].content.decode("utf-8")

    for internal_identifier in (
        "obj-variables",
        "sec-variables",
        "card-variable",
        "worked-variable",
        "practice-variable",
    ):
        assert internal_identifier not in review_html
        assert internal_identifier not in review_markdown
    for reader_label in (
        "Objective 1: Explain and use Python variables.",
        "Variables",
        "Flashcard 1",
        "Worked example 1",
        "Practice exercise 1",
    ):
        assert reader_label in review_html
        assert reader_label in review_markdown


def test_review_headings_survive_valid_cross_namespace_identifier_collisions() -> None:
    """Exercise IDs must not replace objective or section labels in headings."""

    bundle = valid_bundle()
    objective_id = bundle.course.learning_objectives[0].objective_id
    section = bundle.course.modules[0].sections[0]
    bundle.review_pack.worked_examples[0].exercise_id = objective_id
    bundle.review_pack.practice_exercises[0].exercise_id = section.section_id

    rendered = ArtifactRenderer().render_bundle(bundle)
    review_html = rendered["review_pack_html"].content.decode("utf-8")
    review_markdown = rendered["review_pack_markdown"].content.decode("utf-8")

    expected_objective_heading = "Objective 1: Explain and use Python variables."
    assert expected_objective_heading in review_html
    assert expected_objective_heading in review_markdown
    assert "<dt>Variables</dt>" in review_html
    assert "- **Variables:**" in review_markdown


def test_review_renderers_humanize_unresolved_identifier_shaped_prose() -> None:
    """Free-form review steps must not expose stale IDs or schema field names."""

    bundle = valid_bundle()
    bundle.review_pack.review_sequence = [
        (
            "Review objective `obj-variables` instead of stale lo9, complete "
            "`practice-missing-reference` in section-9-stale-title, then "
            "revisit sec9 and the referenced `section_id` before pe9."
        )
    ]

    rendered = ArtifactRenderer().render_bundle(bundle)
    review_text_by_format = {
        "html": rendered["review_pack_html"].content.decode("utf-8"),
        "markdown": rendered["review_pack_markdown"].content.decode("utf-8"),
    }
    pdf_document = fitz.open(
        stream=rendered["review_pack_pdf"].content,
        filetype="pdf",
    )
    try:
        review_text_by_format["pdf"] = "\n".join(
            page.get_text("text") for page in pdf_document
        )
    finally:
        pdf_document.close()
    review_document = Document(BytesIO(rendered["review_pack_docx"].content))
    review_text_by_format["docx"] = "\n".join(
        paragraph.text for paragraph in review_document.paragraphs
    )

    for review_text in review_text_by_format.values():
        assert "Review Objective 1" in review_text
        assert "Review objective Objective 1" not in review_text
        assert "Practice exercise" in review_text
        assert "practice-missing-reference" not in review_text
        assert "the related section" in review_text
        assert "section-9-stale-title" not in review_text
        assert "section_id" not in review_text
        assert "lo9" not in review_text
        assert "sec9" not in review_text
        assert "pe9" not in review_text


def test_empty_optional_course_collections_do_not_render_empty_sections() -> None:
    """An optional empty list must not leave a misleading blank publication section."""

    bundle = valid_bundle()
    bundle.course.modules[0].misconceptions = []
    bundle.course.modules[0].examples = []

    rendered = ArtifactRenderer().render_bundle(bundle)
    course_html = rendered["course_html"].content.decode("utf-8")
    course_markdown = rendered["course_markdown"].content.decode("utf-8")

    assert "Common misconceptions" not in course_html
    assert "Common misconceptions" not in course_markdown
    assert ">Examples<" not in course_html
    assert "### Examples" not in course_markdown


def test_pdf_and_docx_convert_inline_markdown_to_readable_text() -> None:
    """Portable binary formats must not expose backticks or Markdown link syntax."""

    bundle = valid_bundle()
    bundle.review_pack.review_sequence = ["Run `python --version` before review."]
    rendered = ArtifactRenderer().render_bundle(bundle)

    pdf_document = fitz.open(
        stream=rendered["review_pack_pdf"].content,
        filetype="pdf",
    )
    try:
        review_pdf_text = "\n".join(page.get_text("text") for page in pdf_document)
    finally:
        pdf_document.close()
    course_document = Document(BytesIO(rendered["course_docx"].content))
    course_docx_text = "\n".join(
        paragraph.text for paragraph in course_document.paragraphs
    )

    assert "`python --version`" not in review_pdf_text
    assert "python --version" in review_pdf_text
    assert "[The Python Tutorial](" not in course_docx_text
    assert "The Python Tutorial" in course_docx_text
    assert "https://docs.python.org/3/tutorial/" in course_docx_text


def test_code_blocks_keep_format_native_visual_treatment() -> None:
    """Code should remain recognizable in every publication format."""

    bundle = valid_bundle()
    content_block = bundle.course.modules[0].sections[0].content_blocks[0]
    content_block.kind = "code"
    content_block.text = "student_count = 24\nprint(student_count)"

    rendered = ArtifactRenderer().render_bundle(bundle)
    course_html = rendered["course_html"].content.decode("utf-8")
    course_markdown = rendered["course_markdown"].content.decode("utf-8")
    course_pdf = fitz.open(stream=rendered["course_pdf"].content, filetype="pdf")
    try:
        pdf_fonts = {
            span["font"]
            for page in course_pdf
            for block in page.get_text("dict")["blocks"]
            if "lines" in block
            for line in block["lines"]
            for span in line["spans"]
        }
    finally:
        course_pdf.close()
    course_document = Document(BytesIO(rendered["course_docx"].content))

    assert "<pre><code>student_count = 24\nprint(student_count)</code></pre>" in (
        course_html
    )
    assert "```\nstudent_count = 24\nprint(student_count)\n```" in course_markdown
    assert any("Courier" in font_name for font_name in pdf_fonts)
    assert any(
        paragraph.style is not None
        and paragraph.style.name == "Publication Code"
        and "student_count = 24" in paragraph.text
        for paragraph in course_document.paragraphs
    )


def test_pdf_normalizes_common_typographic_punctuation() -> None:
    """English smart punctuation must remain readable with the built-in PDF font."""

    bundle = valid_bundle()
    bundle.course.modules[0].sections[0].content_blocks[
        0
    ].text = "A client\u2019s resolver\u2014when ready\u2014answers \u201csafely\u201d."

    rendered = ArtifactRenderer().render_bundle(bundle)
    pdf_document = fitz.open(
        stream=rendered["course_pdf"].content,
        filetype="pdf",
    )
    try:
        course_pdf_text = "\n".join(page.get_text("text") for page in pdf_document)
    finally:
        pdf_document.close()

    assert "client's resolver-when ready-answers" in course_pdf_text
    assert '"safely"' in course_pdf_text
    assert "client-s" not in course_pdf_text


def test_long_pdf_content_paginates_without_truncating_the_tail() -> None:
    """Long generated lessons must survive page breaks with accurate folios."""

    bundle = valid_bundle()
    bundle.course.title = (
        "A Practical Field Guide to Reliable Variables, Clear Names, "
        "Thoughtful State, and Maintainable Programs"
    )
    long_lesson = " ".join(f"concept-{concept_number}" for concept_number in range(700))
    bundle.course.modules[0].sections[0].content_blocks[
        0
    ].text = f"{long_lesson} final-layout-marker"

    rendered = ArtifactRenderer().render_bundle(bundle)
    pdf_document = fitz.open(stream=rendered["course_pdf"].content, filetype="pdf")
    try:
        assert len(pdf_document) >= 4
        cover_text = pdf_document[0].get_text("text")
        extracted_text = "\n".join(page.get_text("text") for page in pdf_document)
        assert "Maintainable Programs" in cover_text
        assert "final-layout-marker" in extracted_text
        for page_number, page in enumerate(list(pdf_document)[1:], start=2):
            assert f"{page_number} / {len(pdf_document)}" in page.get_text("text")
    finally:
        pdf_document.close()


def test_summary_label_removes_one_model_supplied_leading_colon() -> None:
    """The deterministic Summary label must not produce a visible double colon."""

    bundle = valid_bundle()
    bundle.course.modules[0].sections[
        0
    ].summary = ": Variables give useful names to values."

    rendered = ArtifactRenderer().render_bundle(bundle)
    course_html = rendered["course_html"].content.decode("utf-8")
    course_markdown = rendered["course_markdown"].content.decode("utf-8")

    assert "Summary:</strong> Variables give" in course_html
    assert "Summary:</strong> :" not in course_html
    assert "**Summary:** Variables give" in course_markdown
    assert "**Summary:** :" not in course_markdown


def test_assessment_renderers_use_singular_point_grammar() -> None:
    """A one-point item should read naturally in every student format."""

    bundle = valid_bundle()
    bundle.assessment.blueprint[0].total_points = 1
    bundle.assessment.items[0].points = 1
    bundle.answer_key.answers[0].rubric[0].points = 1

    rendered = ArtifactRenderer().render_bundle(bundle)
    assessment_html = rendered["assessment_html"].content.decode("utf-8")
    assessment_markdown = rendered["assessment_markdown"].content.decode("utf-8")
    pdf_document = fitz.open(
        stream=rendered["assessment_pdf"].content,
        filetype="pdf",
    )
    try:
        assessment_pdf_text = "\n".join(page.get_text("text") for page in pdf_document)
    finally:
        pdf_document.close()

    for assessment_text in (
        assessment_html,
        assessment_markdown,
        assessment_pdf_text,
    ):
        assert "(1 point)" in assessment_text
        assert "(1 points)" not in assessment_text


def test_rtl_course_uses_document_direction_and_searchable_pdf() -> None:
    """RTL metadata reaches HTML and the deterministic PDF contains real text."""

    bundle = valid_bundle()
    bundle.course.language = "he"
    bundle.course.title = (
        "\u05de\u05d1\u05d5\u05d0 \u05dc\u05de\u05e9\u05ea\u05e0\u05d9\u05dd"
    )
    rendered = ArtifactRenderer().render_bundle(bundle)

    course_html = rendered["course_html"].content.decode("utf-8")
    course_pdf = rendered["course_pdf"].content
    pdf_document = fitz.open(stream=course_pdf, filetype="pdf")
    try:
        extracted_pdf_text = "\n".join(page.get_text("text") for page in pdf_document)
    finally:
        pdf_document.close()

    assert '<html lang="he" dir="rtl">' in course_html
    assert course_pdf.startswith(b"%PDF")
    assert "Variables" in extracted_pdf_text


@pytest.mark.parametrize(
    "unsafe_html",
    [
        '<main id="course-content"><script>alert(1)</script></main>',
        '<main id="course-content"><iframe src="https://evil.test"></iframe></main>',
        '<main id="course-content"><img src="https://evil.test/pixel"></main>',
        '<main id="course-content"><p onclick="steal()">Click</p></main>',
        '<main id="course-content">Bearer secret-token /home/ada/private</main>',
    ],
)
def test_rendered_output_qa_rejects_active_or_private_content(
    unsafe_html: str,
) -> None:
    """Post-render QA catches structure-independent safety/privacy failures."""

    with pytest.raises(RenderedOutputQaError):
        validate_rendered_html(
            unsafe_html,
            required_element_ids={"course-content"},
        )


def test_rendered_output_qa_rejects_missing_required_sections() -> None:
    """A visually plausible but incomplete document cannot be delivered."""

    with pytest.raises(RenderedOutputQaError, match="bibliography"):
        validate_rendered_html(
            '<main id="course-content">Course</main>',
            required_element_ids={"course-content", "bibliography"},
        )
